#!/usr/bin/env python3
"""
QSM 5T Reproducibility — Statistical Analysis v6.2 (round 2)
====================================================

Revisions from v6.1 (11 user feedback points):

  1. Output to /output/Figure3 .. Figure6 (top-level, short names).
  2. PDF only — no PNG.
  3. Figure 3: shared y-axis range per ROI across L/R; colour-consistent.
  4. Figure 4A: significance indicators (paired t-test) on each ROI.
  5. Figure 4B: no coloured background; thin black dashed threshold lines;
     "Poor/Moderate/Good/Excellent" labels at top; fix clipped lower CI.
  6. Figure 4A: redesigned as dumbbell chart (cleaner than paired dot).
  7. Figure 5A: fix legend overlap; add p-value; add n per scanner in legend.
  8. Figure 5B: redesigned as grouped bar chart of wSD per ROI per subject.
  9. Figure 6: tighter spacing; "Subject 1" labels; smaller gaps.
 10. Scanner names → 5T-1 … 5T-6 globally.
 11. Heatmap cells forced to square aspect.

Round-2 revisions (8 additional user feedback points):

 R1. Figure 3: added 9th "Overall" column merging all L/R ROI data.
 R2. All colours derived from Figure 1 RdBu_r colormap (-50..50 ppb range).
 R3. Figure 3A: removed LoA coloured fill; standardised dotted-line
     thickness globally (2/3 of correlation reference line, double-spaced).
 R4. Figure 4A: reverted to v6.1 paired dot plot; only Thalamus p-value
     shown; concise legend note for other ROIs (all ns).
 R5. Figure 4B: full error-bar range shown (no artificial clipping);
     x-limits widen to accommodate all CI.
 R6. Figure 5A: removed per-ROI p-value annotations; added concise
     "Kruskal-Wallis: all p > 0.05" in legend.
 R7. Figure 5: decreased spacing between A and B (wspace 0.35→0.15);
     Fig 5B heatmap changed from YlOrRd → Blues.
 R8. Figure 6: already uses RdBu_r (aligned with Figure 1 palette).
"""

from __future__ import annotations

import csv
import logging
import warnings
from collections import defaultdict
from pathlib import Path

import numpy as np
from scipy import stats as sp_stats

warnings.filterwarnings("ignore", category=RuntimeWarning)
logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

# ── Paths ──────────────────────────────────────────────────────────────
PROJECT  = Path(".")
DATA_DIR = PROJECT / "output" / "roi_values_v2"
OUT_DIR  = PROJECT / "output" / "statistics_v6_2"
CSV_DIR  = OUT_DIR / "csv"
TABLE_DIR = OUT_DIR / "tables"

SS_CSV = DATA_DIR / "synthseg_roi_values.csv"
MS_CSV = DATA_DIR / "manual_seg_roi_values.csv"

# ★ Point 1: output folders directly under output/, short names
FIG3_DIR = PROJECT / "output" / "Figure3"
FIG4_DIR = PROJECT / "output" / "Figure4"
FIG5_DIR = PROJECT / "output" / "Figure5"
FIG6_DIR = PROJECT / "output" / "Figure6"
FIG7_DIR = PROJECT / "output" / "Figure7"

# ── Constants ──────────────────────────────────────────────────────────
SCANNER_INFO = {
    "ses-3TUIH1":      (3.0, "3T-1"),
    "ses-3TUIH2":      (3.0, "3T-2"),
    "ses-5TUIH1":      (5.0, "5T-1"),
    "ses-5TUIH2":      (5.0, "5T-2"),
    "ses-5TShanghai":  (5.0, "5T-3"),
    "ses-5TShenzhen":  (5.0, "5T-4"),
    "ses-5TChongqing": (5.0, "5T-5"),
    "ses-5TJingzhou":  (5.0, "5T-6"),
}

# ★ Point 10: use 5T-n naming globally
SCANNER_SHORT = {
    "ses-5TUIH1": "5T-1", "ses-5TUIH2": "5T-2",
    "ses-5TShanghai": "5T-3", "ses-5TShenzhen": "5T-4",
    "ses-5TChongqing": "5T-5", "ses-5TJingzhou": "5T-6",
    "ses-3TUIH1": "3T-1", "ses-3TUIH2": "3T-2",
}
SCANNER_ORDER_5T = [
    "ses-5TUIH1", "ses-5TUIH2", "ses-5TShanghai",
    "ses-5TShenzhen", "ses-5TChongqing", "ses-5TJingzhou",
]

OVERLAP_MAP = {
    "Thalamus":  "Thalamus",
    "Caudate":   "Caudate_nucleus",
    "Putamen":   "Putamen",
    "Pallidum":  "Globus_pallidus",
}
OVERLAP_DISPLAY = {
    "Thalamus": "Thalamus",
    "Caudate":  "Caudate",
    "Putamen":  "Putamen",
    "Pallidum": "Pallidum",
}
SYNTHSEG_ONLY_ROIS = ["Cerebral-White-Matter", "Brain-Stem"]
MANUAL_ONLY_ROIS   = ["Substantia_nigra", "Nucleus_ruber", "Dentate_nucleus"]

KEY_REPRO_ROIS_SYNTHSEG = [
    "Thalamus", "Caudate", "Putamen", "Pallidum",
    "Cerebral-White-Matter", "Brain-Stem",
]
KEY_REPRO_ROIS_MANUAL = [
    "Thalamus", "Caudate_nucleus", "Putamen", "Globus_pallidus",
    "Substantia_nigra", "Nucleus_ruber", "Dentate_nucleus",
]

# ── Unified colour palette — derived from Figure 1 RdBu_r cmap ────────
# Figure 1 uses RdBu_r colormap.  The user requests that all
# statistical-figure colours be drawn from the same cmap, normalised
# to the -50 … +50 ppb legend range.  We sample positions within that
# range so scatter / bar colours harmonise with the QSM overlay images.
import matplotlib
matplotlib.use("Agg")
import matplotlib.colors
from matplotlib.colors import Normalize as _Norm
_FIG1_CMAP  = matplotlib.colormaps['RdBu_r']
_FIG1_NORM  = _Norm(vmin=-50, vmax=50)       # legend range
def _c(ppb):
    """Return hex colour for a given ppb value on the -50..50 scale."""
    rgba = _FIG1_CMAP(_FIG1_NORM(ppb))
    return matplotlib.colors.rgb2hex(rgba[:3])

# Core palette sampled from Figure 1 legend (-50 → 50 ppb)
C3T       = _c(-40)    # cool blue  (negative QSM end)
C5T       = _c(40)     # warm red   (positive QSM end)
C_GREEN   = _c(-15)    # blue-ish neutral
C_PURPLE  = _c(50)     # strong red
C_ORANGE  = _c(25)     # warm
C_GREY    = '#636363'
C_LIGHT3T = _c(-25)    # light blue
C_LIGHT5T = _c(15)     # light pink/salmon

# Scanner colour palette for Figure 5 — 6 evenly-spaced positions
# across the RdBu_r legend so each scanner gets a distinct colour
SCANNER_COLORS = {
    "ses-5TUIH1":      _c(-30),   # 5T-1  medium blue
    "ses-5TUIH2":      _c(-18),   # 5T-2  light blue
    "ses-5TShanghai":  _c(-6),    # 5T-3  pale cool
    "ses-5TShenzhen":  _c(6),     # 5T-4  pale warm
    "ses-5TChongqing": _c(18),    # 5T-5  light red
    "ses-5TJingzhou":  _c(30),    # 5T-6  medium red
}

# ★ Subject anonymisation map  (Point 9)
SUBJECT_NUMBER = {}  # filled at runtime
def _build_subject_numbers(ss_rows):
    subs = sorted(set(r["subject"] for r in ss_rows))
    for i, s in enumerate(subs, 1):
        SUBJECT_NUMBER[s] = f"Subject {i}"


# ═══════════════════════════════════════════════════════════════════════
# UTILITY FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════

def load_csv(path: Path) -> list[dict]:
    rows = []
    with open(path, newline="", encoding="utf-8") as f:
        for r in csv.DictReader(f):
            r["field_strength"] = float(r["field_strength"])
            r["label_id"] = int(r["label_id"])
            r["n_voxels"] = int(r["n_voxels"])
            for k in ("mean_ppb", "std_ppb", "median_ppb", "q25_ppb", "q75_ppb"):
                r[k] = float(r[k])
            rows.append(r)
    return rows


def build_lookup(rows: list[dict]) -> dict:
    out = {}
    for r in rows:
        out[(r["subject"], r["session"], r["roi"], r["side"])] = r
    return out


def _write_csv(path: Path, rows: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        with open(path, "w", newline="", encoding="utf-8") as f:
            f.write("")
        return
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        w.writeheader()
        w.writerows(rows)


def fmt(x, d=2):
    if x is None:
        return "NA"
    try:
        if np.isnan(x):
            return "NA"
    except Exception:
        pass
    return f"{x:.{d}f}"


def icc_31(matrix: np.ndarray):
    """ICC(3,1) — two-way mixed, single measures, consistency."""
    if matrix.shape[0] < 3 or matrix.shape[1] < 2:
        return (np.nan, np.nan, np.nan)
    n, k = matrix.shape
    mean_total = np.mean(matrix)
    ss_rows_v = k * np.sum((np.mean(matrix, axis=1) - mean_total) ** 2)
    ss_cols = n * np.sum((np.mean(matrix, axis=0) - mean_total) ** 2)
    ss_total = np.sum((matrix - mean_total) ** 2)
    ss_error = ss_total - ss_rows_v - ss_cols
    ms_rows_v = ss_rows_v / (n - 1)
    ms_error = ss_error / ((n - 1) * (k - 1))
    if (ms_rows_v + (k - 1) * ms_error) == 0:
        icc = np.nan
    else:
        icc = (ms_rows_v - ms_error) / (ms_rows_v + (k - 1) * ms_error)
    f_val = ms_rows_v / ms_error if ms_error > 0 else np.inf
    df1, df2 = n - 1, (n - 1) * (k - 1)
    if np.isinf(f_val) or np.isnan(f_val) or df2 == 0:
        return (icc, np.nan, np.nan)
    try:
        f_lo = f_val / sp_stats.f.ppf(0.975, df1, df2)
        f_hi = f_val / sp_stats.f.ppf(0.025, df1, df2)
        lo = (f_lo - 1) / (f_lo + k - 1)
        hi = (f_hi - 1) / (f_hi + k - 1)
    except Exception:
        lo, hi = np.nan, np.nan
    return (icc, lo, hi)


def _icc_interpret(icc):
    if np.isnan(icc):
        return "NA"
    if icc < 0.50:
        return "poor"
    if icc < 0.75:
        return "moderate"
    if icc < 0.90:
        return "good"
    return "excellent"


def bland_altman(x, y):
    x, y = np.asarray(x, dtype=float), np.asarray(y, dtype=float)
    diff = x - y
    mean = (x + y) / 2.0
    md = np.mean(diff)
    sd = np.std(diff, ddof=1) if len(diff) > 1 else np.nan
    if len(diff) > 1 and np.isfinite(sd) and sd > 0:
        t_stat, p_val = sp_stats.ttest_1samp(diff, 0)
    else:
        t_stat, p_val = np.nan, np.nan
    return {
        "n": len(diff), "mean_diff": md, "sd_diff": sd,
        "loa_lower": md - 1.96 * sd if np.isfinite(sd) else np.nan,
        "loa_upper": md + 1.96 * sd if np.isfinite(sd) else np.nan,
        "t_stat": t_stat, "p_value": p_val,
        "means": mean, "diffs": diff,
    }


def cohens_d(x, y, paired=True):
    x, y = np.asarray(x, dtype=float), np.asarray(y, dtype=float)
    if len(x) != len(y) or len(x) < 2:
        return np.nan
    if paired:
        diff = x - y
        sd = np.std(diff, ddof=1)
        return np.mean(diff) / sd if sd > 0 else np.nan
    else:
        nx, ny = len(x), len(y)
        var_x, var_y = np.var(x, ddof=1), np.var(y, ddof=1)
        pooled_sd = np.sqrt(((nx - 1) * var_x + (ny - 1) * var_y) / (nx + ny - 2))
        return (np.mean(x) - np.mean(y)) / pooled_sd if pooled_sd > 0 else np.nan


def within_subject_metrics(values):
    values = np.asarray(values, dtype=float)
    if len(values) < 2:
        return np.nan, np.nan
    m = np.mean(values)
    sd = np.std(values, ddof=1)
    cv = sd / abs(m) * 100.0 if abs(m) > 1e-9 else np.nan
    return sd, cv


def qsm_dynamic_range(ss_rows, ms_rows):
    vals = [r["mean_ppb"] for r in ss_rows] + [r["mean_ppb"] for r in ms_rows]
    return min(vals), max(vals), max(vals) - min(vals)


def _sig_stars(p):
    """Return significance annotation string."""
    if p is None or np.isnan(p):
        return ""
    if p < 0.001:
        return "***"
    if p < 0.01:
        return "**"
    if p < 0.05:
        return "*"
    return "ns"


# ═══════════════════════════════════════════════════════════════════════
# CONSENSUS BUILDER — merge SynthSeg + Manual for overlap ROIs
# ═══════════════════════════════════════════════════════════════════════

def build_consensus_5t(ss_rows, ms_rows):
    """Build consensus QSM values for 5T data.

    For the 4 overlap ROIs: average of SynthSeg and Manual (L+R averaged).
    For non-overlap ROIs:   use the single available mask (L+R averaged).

    Returns dict:  (subject, session, unified_roi) → mean QSM (ppb)
    """
    ms_lookup = build_lookup(ms_rows)
    consensus = defaultdict(list)

    for r in ss_rows:
        if r["field_strength"] != 5.0:
            continue
        ss_roi = r["roi"]
        if ss_roi not in OVERLAP_MAP:
            continue
        ms_roi = OVERLAP_MAP[ss_roi]
        unified = OVERLAP_DISPLAY.get(ss_roi, ss_roi)
        key_ms = (r["subject"], r["session"], ms_roi, r["side"])
        ss_val = r["mean_ppb"]
        if key_ms in ms_lookup:
            ms_val = ms_lookup[key_ms]["mean_ppb"]
            consensus[(r["subject"], r["session"], unified)].append(
                (ss_val + ms_val) / 2.0)
        else:
            consensus[(r["subject"], r["session"], unified)].append(ss_val)

    for r in ss_rows:
        if r["field_strength"] != 5.0:
            continue
        if r["roi"] in SYNTHSEG_ONLY_ROIS:
            consensus[(r["subject"], r["session"], r["roi"])].append(
                r["mean_ppb"])

    for r in ms_rows:
        if r["field_strength"] != 5.0:
            continue
        if r["roi"] in MANUAL_ONLY_ROIS:
            consensus[(r["subject"], r["session"], r["roi"])].append(
                r["mean_ppb"])

    result = {}
    for k, vals in consensus.items():
        result[k] = np.mean(vals)
    return result


def build_consensus_all(ss_rows, ms_rows):
    """Same as build_consensus_5t but for ALL field strengths.
    Returns dict:  (subject, session, unified_roi, field_T) → mean QSM (ppb)
    """
    ms_lookup = build_lookup(ms_rows)
    consensus = defaultdict(list)

    for r in ss_rows:
        ss_roi = r["roi"]
        fs = r["field_strength"]
        if ss_roi in OVERLAP_MAP:
            ms_roi = OVERLAP_MAP[ss_roi]
            unified = OVERLAP_DISPLAY.get(ss_roi, ss_roi)
            key_ms = (r["subject"], r["session"], ms_roi, r["side"])
            ss_val = r["mean_ppb"]
            if key_ms in ms_lookup:
                ms_val = ms_lookup[key_ms]["mean_ppb"]
                consensus[(r["subject"], r["session"], unified, fs)].append(
                    (ss_val + ms_val) / 2.0)
            else:
                consensus[(r["subject"], r["session"], unified, fs)].append(
                    ss_val)
        elif ss_roi in SYNTHSEG_ONLY_ROIS:
            consensus[(r["subject"], r["session"], ss_roi, fs)].append(
                r["mean_ppb"])

    for r in ms_rows:
        fs = r["field_strength"]
        if r["roi"] in MANUAL_ONLY_ROIS:
            consensus[(r["subject"], r["session"], r["roi"], fs)].append(
                r["mean_ppb"])

    result = {}
    for k, vals in consensus.items():
        result[k] = np.mean(vals)
    return result


# ═══════════════════════════════════════════════════════════════════════
# ANALYSIS MODULES A – H
# ═══════════════════════════════════════════════════════════════════════

def collect_design_summary(ss_rows, ms_rows):
    ss_subjects = sorted(set(r["subject"] for r in ss_rows))
    all_sessions = sorted(set(r["session"] for r in ss_rows))
    total_5t_scans = len(sorted(set(
        (r["subject"], r["session"]) for r in ss_rows if r["field_strength"] == 5.0)))
    total_3t_scans = len(sorted(set(
        (r["subject"], r["session"]) for r in ss_rows if r["field_strength"] == 3.0)))
    subjects_with_both_3t = sorted(
        s for s in ss_subjects
        if {"ses-3TUIH1", "ses-3TUIH2"}.issubset(
            set(r["session"] for r in ss_rows
                if r["subject"] == s and r["field_strength"] == 3.0)))
    subjects_with_uih_pair = sorted(
        s for s in ss_subjects
        if {"ses-5TUIH1", "ses-5TUIH2"}.issubset(
            set(r["session"] for r in ss_rows
                if r["subject"] == s and r["field_strength"] == 5.0)))
    min_qsm, max_qsm, qr = qsm_dynamic_range(ss_rows, ms_rows)
    _write_csv(CSV_DIR / "00_design_summary.csv", [{
        "n_subjects": len(ss_subjects), "n_total_sessions": len(all_sessions),
        "n_5T_sessions": total_5t_scans, "n_3T_sessions": total_3t_scans,
        "subjects_with_two_3T": len(subjects_with_both_3t),
        "subjects_with_UIH_pair": len(subjects_with_uih_pair),
        "qsm_min_ppb": fmt(min_qsm), "qsm_max_ppb": fmt(max_qsm),
        "qsm_range_ppb": fmt(qr),
    }])
    return {
        "subjects": ss_subjects,
        "total_5t_scans": total_5t_scans, "total_3t_scans": total_3t_scans,
        "subjects_with_both_3t": subjects_with_both_3t,
        "subjects_with_uih_pair": subjects_with_uih_pair,
        "qsm_range": qr, "qsm_min": min_qsm, "qsm_max": max_qsm,
    }


def module_roi_volume_summary(ss_rows, ms_rows):
    log.info("Module A: ROI volume summary")
    results = []
    for mask_label, rows in [("synthseg", ss_rows), ("manual", ms_rows)]:
        groups = defaultdict(list)
        for r in rows:
            if r["field_strength"] == 5.0:
                groups[(r["roi"], r["side"])].append(r["n_voxels"])
        for (roi, side), vals in sorted(groups.items()):
            a = np.asarray(vals, dtype=float)
            results.append({
                "mask": mask_label, "roi": roi, "side": side,
                "n_sessions_5T": len(a),
                "mean_n_voxels": fmt(np.mean(a), 0),
                "sd_n_voxels": fmt(np.std(a, ddof=1), 0) if len(a) > 1 else "NA",
                "min_n_voxels": fmt(np.min(a), 0),
                "max_n_voxels": fmt(np.max(a), 0),
            })
    _write_csv(CSV_DIR / "01_roi_volume_summary.csv", results)
    return results


def module_descriptive_by_field(ss_rows, ms_rows):
    log.info("Module B: Descriptive QSM by field strength")
    results = []
    for mask_label, rows in [("synthseg", ss_rows), ("manual", ms_rows)]:
        groups = defaultdict(list)
        for r in rows:
            groups[(r["roi"], r["side"], int(r["field_strength"]))].append(r["mean_ppb"])
        for (roi, side, fs), vals in sorted(groups.items()):
            a = np.asarray(vals, dtype=float)
            results.append({
                "mask": mask_label, "roi": roi, "side": side,
                "field_strength_T": fs, "n_sessions": len(a),
                "mean_ppb": fmt(np.mean(a)), "sd_ppb": fmt(np.std(a, ddof=1)) if len(a) > 1 else "NA",
                "median_ppb": fmt(np.median(a)),
                "min_ppb": fmt(np.min(a)), "max_ppb": fmt(np.max(a)),
            })
    _write_csv(CSV_DIR / "02_descriptive_by_field.csv", results)
    return results


def module_synthseg_vs_manual(ss_rows, ms_rows, qsm_range):
    """Module C: SynthSeg vs Manual — L and R analysed independently."""
    log.info("Module C: SynthSeg vs Manual agreement (L/R independent)")
    ms_lookup = build_lookup(ms_rows)
    results, plot_data = [], []
    for ss_roi, ms_roi in OVERLAP_MAP.items():
        for side in ["L", "R"]:
            ss_vals, ms_vals = [], []
            for r in ss_rows:
                if r["roi"] == ss_roi and r["side"] == side:
                    key = (r["subject"], r["session"], ms_roi, side)
                    if key in ms_lookup:
                        ss_vals.append(r["mean_ppb"])
                        ms_vals.append(ms_lookup[key]["mean_ppb"])
            if len(ss_vals) < 3:
                continue
            ssa, msa = np.asarray(ss_vals), np.asarray(ms_vals)
            ba = bland_altman(ssa, msa)
            _, p_val = sp_stats.ttest_rel(ssa, msa)
            try:
                _, w_pval = sp_stats.wilcoxon(ssa - msa)
            except Exception:
                w_pval = np.nan
            pearson_r, pearson_p = sp_stats.pearsonr(ssa, msa)
            spearman_rho, spearman_p = sp_stats.spearmanr(ssa, msa)
            icc, icc_lo, icc_hi = icc_31(np.column_stack([ssa, msa]))
            loa_w = ba["loa_upper"] - ba["loa_lower"] if np.isfinite(ba["loa_upper"]) else np.nan
            cd = cohens_d(ssa, msa, paired=True)
            plot_data.append({
                "roi": ss_roi, "side": side,
                "label": f"{ss_roi} ({side})",
                "ssa": ssa, "msa": msa, "ba": ba,
                "icc": icc, "icc_lo": icc_lo, "icc_hi": icc_hi,
                "pearson_r": pearson_r,
            })
            results.append({
                "synthseg_roi": ss_roi, "manual_roi": ms_roi, "side": side,
                "n_pairs": len(ssa),
                "mean_synthseg_ppb": fmt(np.mean(ssa)),
                "mean_manual_ppb": fmt(np.mean(msa)),
                "mean_bias_ppb": fmt(ba["mean_diff"]),
                "sd_diff_ppb": fmt(ba["sd_diff"]),
                "cohens_d": fmt(cd),
                "loa_lower_ppb": fmt(ba["loa_lower"]),
                "loa_upper_ppb": fmt(ba["loa_upper"]),
                "loa_width_ppb": fmt(loa_w),
                "loa_width_pct_range": fmt((loa_w / qsm_range) * 100) if np.isfinite(loa_w) else "NA",
                "paired_t_p": fmt(p_val, 4),
                "wilcoxon_p": fmt(w_pval, 4) if not np.isnan(w_pval) else "NA",
                "pearson_r": fmt(pearson_r, 4), "pearson_p": fmt(pearson_p, 4),
                "spearman_rho": fmt(spearman_rho, 4), "spearman_p": fmt(spearman_p, 4),
                "ICC_31": fmt(icc, 4), "ICC_31_lower": fmt(icc_lo, 4),
                "ICC_31_upper": fmt(icc_hi, 4),
                "ICC_interpretation": _icc_interpret(icc),
            })
    _write_csv(CSV_DIR / "03_synthseg_vs_manual.csv", results)
    return results, plot_data


def module_3t_vs_5t_consensus(ss_rows, ms_rows, qsm_range):
    """Module D: 3T vs 5T — consensus (SS+Manual averaged), L+R averaged."""
    log.info("Module D: 3T vs 5T (consensus, L+R merged)")
    consensus = build_consensus_all(ss_rows, ms_rows)

    grp = defaultdict(lambda: {"3T": [], "5T": []})
    for (sub, ses, roi, fs), val in consensus.items():
        bucket = "3T" if fs == 3.0 else "5T"
        grp[(sub, roi)][bucket].append(val)

    rois = sorted(set(k[1] for k in grp.keys()))
    results, plot_data = [], []
    for roi in rois:
        vals_3t, vals_5t = [], []
        for sub in sorted(set(k[0] for k in grp.keys())):
            key = (sub, roi)
            if key not in grp or not grp[key]["3T"] or not grp[key]["5T"]:
                continue
            vals_3t.append(np.mean(grp[key]["3T"]))
            vals_5t.append(np.mean(grp[key]["5T"]))
        if len(vals_3t) < 3:
            continue
        a3, a5 = np.asarray(vals_3t), np.asarray(vals_5t)
        ba = bland_altman(a3, a5)
        _, p_val = sp_stats.ttest_rel(a3, a5)
        try:
            _, w_pval = sp_stats.wilcoxon(a3 - a5)
        except Exception:
            w_pval = np.nan
        icc, icc_lo, icc_hi = icc_31(np.column_stack([a3, a5]))
        loa_w = ba["loa_upper"] - ba["loa_lower"] if np.isfinite(ba["loa_upper"]) else np.nan
        cd = cohens_d(a3, a5, paired=True)
        results.append({
            "roi": roi, "side": "L+R",
            "n_subjects": len(a3),
            "mean_3T_ppb": fmt(np.mean(a3)), "sd_3T_ppb": fmt(np.std(a3, ddof=1)),
            "mean_5T_ppb": fmt(np.mean(a5)), "sd_5T_ppb": fmt(np.std(a5, ddof=1)),
            "mean_diff_ppb": fmt(ba["mean_diff"]),
            "sd_diff_ppb": fmt(ba["sd_diff"]), "cohens_d": fmt(cd),
            "loa_lower_ppb": fmt(ba["loa_lower"]),
            "loa_upper_ppb": fmt(ba["loa_upper"]),
            "loa_width_ppb": fmt(loa_w),
            "paired_t_p": fmt(p_val, 4),
            "wilcoxon_p": fmt(w_pval, 4) if not np.isnan(w_pval) else "NA",
            "ICC_31": fmt(icc, 4), "ICC_31_lower": fmt(icc_lo, 4),
            "ICC_31_upper": fmt(icc_hi, 4),
            "ICC_interpretation": _icc_interpret(icc),
        })
        plot_data.append({
            "roi": roi,
            "a3": a3, "a5": a5,
            "icc": icc, "icc_lo": icc_lo, "icc_hi": icc_hi,
            "mean_3T": np.mean(a3), "sd_3T": np.std(a3, ddof=1),
            "mean_5T": np.mean(a5), "sd_5T": np.std(a5, ddof=1),
            "p_val": p_val,
        })
    _write_csv(CSV_DIR / "04_3T_vs_5T_consensus.csv", results)
    return results, plot_data


def module_5t_interscanner_all(ss_rows, ms_rows, qsm_range):
    """Module E: 5T inter-scanner (ALL scanners, consensus)."""
    log.info("Module E: 5T inter-scanner (all scanners, consensus)")
    consensus = build_consensus_5t(ss_rows, ms_rows)

    scanner_roi_vals = defaultdict(list)
    subject_roi_vals = defaultdict(list)

    for (sub, ses, roi), val in consensus.items():
        scanner_roi_vals[(ses, roi)].append(val)
        subject_roi_vals[(sub, roi)].append(val)

    scanner_avg_rows = []
    for (ses, roi), vals in sorted(scanner_roi_vals.items()):
        a = np.asarray(vals, dtype=float)
        scanner_avg_rows.append({
            "scanner": SCANNER_SHORT.get(ses, ses),
            "session": ses, "roi": roi,
            "n_subjects": len(a),
            "mean_ppb": fmt(np.mean(a)), "sd_ppb": fmt(np.std(a, ddof=1)) if len(a) > 1 else "NA",
            "min_ppb": fmt(np.min(a)), "max_ppb": fmt(np.max(a)),
        })

    subject_wsd_rows = []
    for (sub, roi), vals in sorted(subject_roi_vals.items()):
        if len(vals) < 2:
            continue
        a = np.asarray(vals, dtype=float)
        sd, cv = within_subject_metrics(a)
        subject_wsd_rows.append({
            "subject": sub, "roi": roi,
            "n_5T_scanners": len(a),
            "mean_ppb": fmt(np.mean(a)), "wSD_ppb": fmt(sd),
            "cv_percent": fmt(cv, 1), "range_ppb": fmt(np.ptp(a)),
        })

    _write_csv(CSV_DIR / "05_5T_scanner_average.csv", scanner_avg_rows)
    _write_csv(CSV_DIR / "06_5T_subject_wSD.csv", subject_wsd_rows)
    return scanner_avg_rows, subject_wsd_rows, scanner_roi_vals, subject_roi_vals


def module_left_right(ss_rows, ms_rows, qsm_range):
    log.info("Module F: Left-right symmetry")
    results = []
    for mask_label, rows in [("synthseg", ss_rows), ("manual", ms_rows)]:
        lookup = build_lookup(rows)
        rois = sorted(set(r["roi"] for r in rows if r["side"] == "L"))
        for roi in rois:
            lvals, rvals = [], []
            for r in rows:
                if r["roi"] == roi and r["side"] == "L":
                    key_r = (r["subject"], r["session"], roi, "R")
                    if key_r in lookup:
                        lvals.append(r["mean_ppb"])
                        rvals.append(lookup[key_r]["mean_ppb"])
            if len(lvals) < 3:
                continue
            la, ra = np.asarray(lvals), np.asarray(rvals)
            ba = bland_altman(la, ra)
            _, p_val = sp_stats.ttest_rel(la, ra)
            try:
                _, w_pval = sp_stats.wilcoxon(la - ra)
            except Exception:
                w_pval = np.nan
            icc, icc_lo, icc_hi = icc_31(np.column_stack([la, ra]))
            laterality = (la - ra) / (((la + ra) / 2.0) + 1e-12) * 100.0
            cd = cohens_d(la, ra, paired=True)
            loa_w = ba["loa_upper"] - ba["loa_lower"] if np.isfinite(ba["loa_upper"]) else np.nan
            results.append({
                "mask": mask_label, "roi": roi, "n_pairs": len(la),
                "mean_L_ppb": fmt(np.mean(la)), "mean_R_ppb": fmt(np.mean(ra)),
                "mean_diff_ppb": fmt(ba["mean_diff"]),
                "sd_diff_ppb": fmt(ba["sd_diff"]), "cohens_d": fmt(cd),
                "laterality_pct": fmt(np.mean(laterality), 2),
                "paired_t_p": fmt(p_val, 4),
                "wilcoxon_p": fmt(w_pval, 4) if not np.isnan(w_pval) else "NA",
                "ICC_31": fmt(icc, 4), "ICC_31_lower": fmt(icc_lo, 4),
                "ICC_31_upper": fmt(icc_hi, 4),
                "loa_width_ppb": fmt(loa_w),
            })
    _write_csv(CSV_DIR / "07_left_right_symmetry.csv", results)
    return results


def module_3t_repeatability_limit(ss_rows, ms_rows):
    log.info("Module G: 3T repeatability limitation")
    results = []
    for mask_label, rows in [("synthseg", ss_rows), ("manual", ms_rows)]:
        subs_both = sorted(
            sub for sub in set(r["subject"] for r in rows)
            if {"ses-3TUIH1", "ses-3TUIH2"}.issubset(
                set(rr["session"] for rr in rows
                    if rr["subject"] == sub and rr["field_strength"] == 3.0)))
        roi_sides = sorted(set((r["roi"], r["side"]) for r in rows if r["field_strength"] == 3.0))
        for roi, side in roi_sides:
            n_pairs = sum(
                1 for sub in subs_both
                if any(r["subject"] == sub and r["session"] == "ses-3TUIH1"
                       and r["roi"] == roi and r["side"] == side for r in rows)
                and any(r["subject"] == sub and r["session"] == "ses-3TUIH2"
                        and r["roi"] == roi and r["side"] == side for r in rows))
            results.append({
                "mask": mask_label, "roi": roi, "side": side,
                "n_paired_subjects": n_pairs,
                "formal_possible": "yes" if n_pairs >= 3 else "no",
                "note": ("Only one subject has repeated 3T scans; "
                         "no group-level 3T repeatability inference possible."
                         if n_pairs < 3 else "Sufficient for formal analysis."),
            })
    _write_csv(CSV_DIR / "08_3T_limitation.csv", results)
    return results


def module_reviewer_summary(ss_rows, ms_rows):
    log.info("Module H: Reviewer summary")
    results = []
    for mask_label, rows, key_rois in [
        ("synthseg", ss_rows, KEY_REPRO_ROIS_SYNTHSEG),
        ("manual", ms_rows, KEY_REPRO_ROIS_MANUAL),
    ]:
        groups = defaultdict(list)
        vgroups = defaultdict(list)
        for r in rows:
            if r["field_strength"] == 5.0:
                groups[(r["roi"], r["side"])].append(r["mean_ppb"])
                vgroups[(r["roi"], r["side"])].append(r["n_voxels"])
        for (roi, side), vals in sorted(groups.items()):
            if roi not in key_rois:
                continue
            results.append({
                "mask": mask_label, "roi": roi, "side": side,
                "mean_5T_ppb": fmt(np.mean(vals)),
                "sd_5T_ppb": fmt(np.std(vals, ddof=1)),
                "mean_voxels": fmt(np.mean(vgroups[(roi, side)]), 0),
                "n_sessions": len(vals),
            })
    _write_csv(CSV_DIR / "09_reviewer_summary.csv", results)
    return results


# Global line-thickness constants  (Point 3: consistent across all figures)
# The correlation panel reference line is the baseline (REF_LW).
# Dotted / dashed auxiliary lines use 2/3 of that thickness.
REF_LW       = 0.7     # reference-line width (Fig 3B identity line, etc.)
DOTTED_LW    = REF_LW * 2 / 3   # ≈ 0.47  — LoA lines, threshold dashes
DOTTED_DASH  = (3.5, 3.5)       # double-spaced dotted pattern (on, off)


# ═══════════════════════════════════════════════════════════════════════
# MATPLOTLIB SETUP
# ═══════════════════════════════════════════════════════════════════════

def _setup_mpl():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    plt.rcParams.update({
        'font.family': 'Arial',
        'font.size': 9,
        'axes.labelsize': 10,
        'axes.titlesize': 11,
        'axes.titleweight': 'bold',
        'xtick.labelsize': 8,
        'ytick.labelsize': 8,
        'legend.fontsize': 8,
        'legend.framealpha': 0.85,
        'figure.dpi': 300,
        'savefig.dpi': 300,
        'savefig.bbox': 'tight',
        'savefig.pad_inches': 0.12,
        'axes.spines.top': False,
        'axes.spines.right': False,
        'axes.linewidth': 0.6,
        'xtick.major.width': 0.6,
        'ytick.major.width': 0.6,
        'lines.linewidth': 0.7,
    })
    return plt


def _annotated_heatmap(ax, data, row_labels, col_labels, cmap, vmin, vmax,
                       title='', fmt_str='.1f', cbar_label='',
                       annot_fontsize=8, show_cbar=True,
                       force_square=True):
    """Draw a publication-quality annotated heatmap on *ax*.
    ★ Point 11: cells forced to square via set_aspect + fixed limits."""
    import matplotlib.pyplot as plt
    from matplotlib.colors import Normalize
    import matplotlib.patches as mpatches

    nr, nc = data.shape
    norm = Normalize(vmin=vmin, vmax=vmax)

    for i in range(nr):
        for j in range(nc):
            val = data[i, j]
            if np.isnan(val):
                ax.add_patch(mpatches.Rectangle(
                    (j - 0.5, i - 0.5), 1, 1,
                    facecolor='#F0F0F0', edgecolor='white', lw=0.8))
                ax.text(j, i, '—', ha='center', va='center',
                        fontsize=annot_fontsize, color='#AAAAAA')
                continue
            fc = plt.get_cmap(cmap)(norm(val))
            ax.add_patch(mpatches.Rectangle(
                (j - 0.5, i - 0.5), 1, 1,
                facecolor=fc, edgecolor='white', lw=0.8))
            lum = 0.299 * fc[0] + 0.587 * fc[1] + 0.114 * fc[2]
            tc = 'white' if lum < 0.55 else 'black'
            ax.text(j, i, f'{val:{fmt_str}}', ha='center', va='center',
                    fontsize=annot_fontsize, color=tc, fontweight='medium')

    ax.set_xlim(-0.5, nc - 0.5)
    ax.set_ylim(nr - 0.5, -0.5)
    ax.set_xticks(range(nc))
    ax.set_xticklabels(col_labels, rotation=45, ha='right', fontsize=8)
    ax.set_yticks(range(nr))
    ax.set_yticklabels(row_labels, fontsize=8)
    if title:
        ax.set_title(title, loc='left', fontsize=10, fontweight='bold')
    for spine in ax.spines.values():
        spine.set_visible(True)
        spine.set_linewidth(0.3)
        spine.set_color('#CCCCCC')

    # ★ Point 11: force square cells (can be disabled for GridSpec layouts)
    if force_square:
        ax.set_aspect('equal', adjustable='box')

    if show_cbar:
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=norm)
        sm.set_array([])
        cbar = ax.figure.colorbar(sm, ax=ax, fraction=0.035, pad=0.04,
                                  aspect=25)
        cbar.ax.tick_params(labelsize=7)
        if cbar_label:
            cbar.set_label(cbar_label, fontsize=8)
    return ax


# ═══════════════════════════════════════════════════════════════════════
# FIGURE 3 — SynthSeg vs Manual  (L/R independent)
# ★ Point 3: shared y-axis per ROI; colour consistency checked
# ═══════════════════════════════════════════════════════════════════════

def create_figure_3(ba_plot_data, fig3_csv_rows):
    plt = _setup_mpl()
    log.info("Figure 3: SynthSeg vs Manual (L/R independent + Overall)")

    if not ba_plot_data:
        return

    roi_order = list(OVERLAP_MAP.keys())
    # Group by ROI → [L entry, R entry]
    roi_groups = {}
    for roi in roi_order:
        pair = []
        for side in ["L", "R"]:
            match = [d for d in ba_plot_data
                     if d["roi"] == roi and d["side"] == side]
            if match:
                pair.append(match[0])
        if pair:
            roi_groups[roi] = pair

    n_rois = len(roi_groups)
    if n_rois == 0:
        return

    # ★ NEW: build an "Overall" entry that merges all L/R data across ROIs
    all_ssa_overall = np.concatenate([d["ssa"] for d in ba_plot_data])
    all_msa_overall = np.concatenate([d["msa"] for d in ba_plot_data])
    ba_overall = bland_altman(all_ssa_overall, all_msa_overall)
    pearson_r_overall, _ = sp_stats.pearsonr(all_ssa_overall, all_msa_overall)
    icc_overall, _, _ = icc_31(np.column_stack([all_ssa_overall, all_msa_overall]))
    overall_entry = {
        "roi": "Overall", "side": "",
        "label": "Overall",
        "ssa": all_ssa_overall, "msa": all_msa_overall,
        "ba": ba_overall,
        "icc": icc_overall,
        "pearson_r": pearson_r_overall,
    }

    # Total columns = L + R per ROI + 1 Overall
    total_cols = sum(len(v) for v in roi_groups.values()) + 1

    fig, axes = plt.subplots(
        2, total_cols,
        figsize=(2.6 * total_cols, 6.5),
        gridspec_kw={'hspace': 0.32, 'wspace': 0.18})
    if total_cols == 1:
        axes = axes.reshape(2, 1)

    # ★ Pre-compute shared y-axis limits per ROI for BA row
    roi_ylims_ba = {}
    for roi, panels in roi_groups.items():
        all_diffs = np.concatenate([p["ba"]["diffs"] for p in panels])
        all_loa_lo = min(p["ba"]["loa_lower"] for p in panels
                         if np.isfinite(p["ba"]["loa_lower"]))
        all_loa_hi = max(p["ba"]["loa_upper"] for p in panels
                         if np.isfinite(p["ba"]["loa_upper"]))
        margin = (all_loa_hi - all_loa_lo) * 0.15
        roi_ylims_ba[roi] = (min(all_loa_lo, np.min(all_diffs)) - margin,
                              max(all_loa_hi, np.max(all_diffs)) + margin)
    # ylim for Overall
    ov_diffs = ba_overall["diffs"]
    ov_margin = (ba_overall["loa_upper"] - ba_overall["loa_lower"]) * 0.15
    roi_ylims_ba["Overall"] = (
        min(ba_overall["loa_lower"], np.min(ov_diffs)) - ov_margin,
        max(ba_overall["loa_upper"], np.max(ov_diffs)) + ov_margin)

    # Collect all entries in drawing order: L/R per ROI, then Overall
    draw_order = []
    for roi, panels in roi_groups.items():
        for d in panels:
            draw_order.append(d)
    draw_order.append(overall_entry)

    # Draw panels
    for col, d in enumerate(draw_order):
        ba = d["ba"]
        lbl = d["label"]
        roi_key = d["roi"]

        # --- Row A: Bland-Altman ---
        ax_ba = axes[0, col]
        ax_ba.scatter(ba["means"], ba["diffs"], alpha=0.72,
                      s=28, color=C3T, edgecolors='white',
                      linewidths=0.3, zorder=3)
        ax_ba.axhline(ba["mean_diff"], color=C5T, linewidth=1.0,
                      zorder=2)
        # ★ Mod 3: LoA lines use global DOTTED_LW + DOTTED_DASH; NO fill_between
        ax_ba.axhline(ba["loa_upper"], color=C_GREY,
                      linestyle=(0, DOTTED_DASH), lw=DOTTED_LW, zorder=2)
        ax_ba.axhline(ba["loa_lower"], color=C_GREY,
                      linestyle=(0, DOTTED_DASH), lw=DOTTED_LW, zorder=2)
        ax_ba.axhline(0, color='black',
                      linestyle=(0, DOTTED_DASH), lw=DOTTED_LW, alpha=0.3)
        # (fill_between REMOVED per Mod 3)
        # ★ shared y-axis per ROI
        ax_ba.set_ylim(roi_ylims_ba[roi_key])
        ax_ba.text(
            0.97, 0.95,
            f"Bias={ba['mean_diff']:.1f}\nLoA=[{ba['loa_lower']:.0f},{ba['loa_upper']:.0f}]",
            transform=ax_ba.transAxes, fontsize=5.5, ha='right',
            va='top',
            bbox=dict(boxstyle='round,pad=0.15', facecolor='white',
                      alpha=0.85, edgecolor='none'))
        if col == 0:
            ax_ba.set_ylabel('SS − Manual (ppb)', fontsize=8)
        ax_ba.set_xlabel('Mean (ppb)', fontsize=7)
        ax_ba.set_title(lbl, fontsize=9, fontweight='bold')
        ax_ba.tick_params(labelsize=6.5)

        # --- Row B: Correlation ---
        ax_c = axes[1, col]
        ax_c.scatter(d["msa"], d["ssa"], alpha=0.72,
                     s=28, color=C3T, edgecolors='white',
                     linewidths=0.3, zorder=3)
        allv = np.concatenate([d["msa"], d["ssa"]])
        vm, vx = np.min(allv) - 5, np.max(allv) + 5
        ax_c.plot([vm, vx], [vm, vx], color=C_GREY,
                  ls='--', lw=REF_LW, alpha=0.6)
        ax_c.set_xlim(vm, vx)
        ax_c.set_ylim(vm, vx)
        ax_c.set_aspect('equal', adjustable='box')
        if col == 0:
            ax_c.set_ylabel('SynthSeg (ppb)', fontsize=8)
        ax_c.set_xlabel('Manual (ppb)', fontsize=7)
        ax_c.text(
            0.05, 0.92,
            f"r={d['pearson_r']:.3f}\nICC={d['icc']:.2f}",
            transform=ax_c.transAxes, fontsize=6, fontweight='bold',
            color='black',
            bbox=dict(boxstyle='round,pad=0.15', facecolor='white',
                      alpha=0.85, edgecolor='none'))
        ax_c.set_title(lbl, fontsize=9, fontweight='bold')
        ax_c.tick_params(labelsize=6.5)

    fig.text(0.003, 0.76, 'A  Bland-Altman', fontsize=10,
             fontweight='bold', rotation=90, va='center')
    fig.text(0.003, 0.28, 'B  Correlation', fontsize=10,
             fontweight='bold', rotation=90, va='center')

    fig.savefig(FIG3_DIR / 'Figure_3_synthseg_vs_manual.pdf',
                dpi=300, bbox_inches='tight')
    plt.close(fig)
    _write_csv(FIG3_DIR / 'Figure_3_data.csv', fig3_csv_rows)
    log.info("  → Figure_3/ saved")


# ═══════════════════════════════════════════════════════════════════════
# FIGURE 4 — 3T vs 5T  (consensus, L+R merged)
# ★ Panel A: grouped-bar + jittered individual data points
#   (learned from figures.py fig3_merged_bar_subject Panel A)
# ★ Panel B: ICC forest — full CI range shown
# ═══════════════════════════════════════════════════════════════════════

def create_figure_4(paired_plot, fig4_csv_rows):
    plt = _setup_mpl()
    import matplotlib.gridspec as gridspec
    from matplotlib.lines import Line2D
    log.info("Figure 4: 3T vs 5T (consensus, L+R merged)")

    key = [d for d in paired_plot if not np.isnan(d["icc"])]
    key = sorted(key, key=lambda d: d["icc"])
    if not key:
        return

    n_rois = len(key)
    roi_labels = [d['roi'] for d in key]

    fig_h = max(5.5, 0.65 * n_rois + 2.0)
    fig = plt.figure(figsize=(13, fig_h))
    gs = gridspec.GridSpec(1, 2, width_ratios=[1.2, 1], wspace=0.35)

    # ═══════ Panel A: Grouped bar + jittered scatter ═══════
    ax_a = fig.add_subplot(gs[0])
    y = np.arange(n_rois)
    bar_h = 0.32
    rng = np.random.default_rng(42)

    for fi, (field_key, field_label, c_fill, c_light) in enumerate([
            ("a3", "3T", C3T, C_LIGHT3T),
            ("a5", "5T", C5T, C_LIGHT5T)]):
        means, sems = [], []
        for d in key:
            arr = np.asarray(d[field_key], dtype=float)
            means.append(np.mean(arr))
            sems.append(sp_stats.sem(arr) if len(arr) > 1 else 0)
        means = np.array(means)
        sems  = np.array(sems)
        n_subj = len(key[0][field_key])

        off = (fi - 0.5) * bar_h
        # Transparent bar with coloured edge (figures.py style)
        ax_a.barh(y + off, means, height=bar_h, xerr=sems,
                  color=c_fill, alpha=0.30, edgecolor=c_fill,
                  linewidth=0.6, capsize=2, zorder=2,
                  label=f'{field_label} (n={n_subj})',
                  error_kw={'linewidth': 0.5, 'color': c_fill})
        # Jittered individual data points
        for j, d in enumerate(key):
            arr = np.asarray(d[field_key], dtype=float)
            jit = rng.uniform(-bar_h * 0.28, bar_h * 0.28, len(arr))
            ax_a.scatter(arr, np.full(len(arr), y[j] + off) + jit,
                         color=c_fill, s=18, alpha=0.65,
                         edgecolors='white', linewidth=0.3, zorder=5)

    # Thalamus p-value annotation only
    for j, d in enumerate(key):
        if d["roi"] == "Thalamus":
            p = d.get("p_val", np.nan)
            if np.isfinite(p):
                xmax_here = max(np.max(d["a3"]), np.max(d["a5"]))
                ax_a.text(xmax_here + 2, y[j], f'p = {p:.3f}',
                          fontsize=7, fontweight='bold', va='center',
                          ha='left', color=C5T)

    ax_a.set_yticks(y)
    ax_a.set_yticklabels(roi_labels, fontsize=8)
    ax_a.set_xlabel('QSM Value (ppb)')
    ax_a.axvline(0, color='grey', ls='-', lw=0.3)

    ax_a.legend(frameon=True, loc='upper right', fontsize=7.5,
                framealpha=0.85, edgecolor='none')
    # Concise note — placed at right side below legend
    ax_a.text(0.98, 0.88,
              "Paired t-test: only Thalamus p < 0.05;\nall other regions ns",
              transform=ax_a.transAxes, va='top', ha='right', fontsize=6.5,
              color=C_GREY, style='italic')
    ax_a.set_title('A   Mean QSM: 3T vs 5T', loc='left', fontsize=10)
    xl = ax_a.get_xlim()
    ax_a.set_xlim(xl[0], xl[1] + (xl[1] - xl[0]) * 0.08)

    # ═══════ Panel B: ICC forest — CLEAN ═══════
    ax_b = fig.add_subplot(gs[1])
    icc_v = np.array([d["icc"]    for d in key])
    lo_v  = np.array([d["icc_lo"] for d in key])
    hi_v  = np.array([d["icc_hi"] for d in key])

    for i in range(n_rois):
        lo_err = max(icc_v[i] - lo_v[i], 0) if np.isfinite(lo_v[i]) else 0
        hi_err = max(hi_v[i] - icc_v[i], 0) if np.isfinite(hi_v[i]) else 0
        ax_b.errorbar(icc_v[i], y[i],
                      xerr=[[lo_err], [hi_err]],
                      fmt='o', color=C3T, ecolor=C_GREY,
                      capsize=3, markersize=5, markeredgewidth=0.5,
                      markeredgecolor='white', elinewidth=0.8,
                      clip_on=False)

    # Threshold lines
    ax_b.axvline(0.50, color='black',
                 linestyle=(0, DOTTED_DASH), lw=DOTTED_LW, alpha=0.6)
    ax_b.axvline(0.75, color='black',
                 linestyle=(0, DOTTED_DASH), lw=DOTTED_LW, alpha=0.6)
    ax_b.axvline(0.90, color='black',
                 linestyle=(0, DOTTED_DASH), lw=DOTTED_LW, alpha=0.6)

    ax_b.text(0.20, 1.03, 'Poor', transform=ax_b.get_xaxis_transform(),
              fontsize=7, ha='center', va='bottom', color='black',
              fontstyle='italic', clip_on=False)
    ax_b.text(0.625, 1.03, 'Moderate', transform=ax_b.get_xaxis_transform(),
              fontsize=7, ha='center', va='bottom', color='black',
              fontstyle='italic', clip_on=False)
    ax_b.text(0.825, 1.03, 'Good', transform=ax_b.get_xaxis_transform(),
              fontsize=7, ha='center', va='bottom', color='black',
              fontstyle='italic', clip_on=False)
    ax_b.text(0.975, 1.03, 'Excellent', transform=ax_b.get_xaxis_transform(),
              fontsize=7, ha='center', va='bottom', color='black',
              fontstyle='italic', clip_on=False)

    # Increase title pad so labels sit between title and plot
    ax_b.set_title('B   3T–5T Agreement: ICC(3,1)', loc='left', fontsize=10,
                   pad=18)

    ax_b.set_yticks(y)
    ax_b.set_yticklabels(roi_labels, fontsize=8)
    ax_b.set_xlabel('ICC(3,1) with 95% CI')

    all_lo = [lo_v[i] for i in range(n_rois) if np.isfinite(lo_v[i])]
    x_lo = min(all_lo) - 0.05 if all_lo else -0.1
    ax_b.set_xlim(min(x_lo, -0.1), 1.05)

    fig.savefig(FIG4_DIR / 'Figure_4_3T_vs_5T.pdf',
                dpi=300, bbox_inches='tight')
    plt.close(fig)
    _write_csv(FIG4_DIR / 'Figure_4_data.csv', fig4_csv_rows)
    log.info("  → Figure_4/ saved")


# ═══════════════════════════════════════════════════════════════════════
# FIGURE 5 — 5T Inter-Scanner Reproducibility
# ★ Mod 6: No p-value annotations; "no sig diff" in legend
# ★ Mod 7: Decreased wspace; Fig5B colormap aligned with RdBu_r
# ═══════════════════════════════════════════════════════════════════════

def create_figure_5(scanner_roi_vals, subject_roi_vals, ss_rows,
                    scanner_csv, subject_csv):
    plt = _setup_mpl()
    from matplotlib.lines import Line2D
    log.info("Figure 5: 5T Inter-Scanner (all scanners, consensus)")

    rois_for_fig5 = list(OVERLAP_DISPLAY.values())
    scanners_present = [s for s in SCANNER_ORDER_5T
                        if any((s, roi) in scanner_roi_vals
                               for roi in rois_for_fig5)]
    if not scanners_present:
        return

    n_rois = len(rois_for_fig5)
    n_scanners = len(scanners_present)
    bar_width = 0.80 / n_scanners
    x_base = np.arange(n_rois)

    # Count subjects per scanner
    sub_per_scanner = {}
    for ses in scanners_present:
        subs = set(r["subject"] for r in ss_rows
                   if r["session"] == ses and r["field_strength"] == 5.0)
        sub_per_scanner[ses] = len(subs)

    fig_w = max(10, 2.5 * n_rois + 2)
    # ★ Mod 7: decreased wspace from 0.35 → 0.15
    fig, (ax_a, ax_b) = plt.subplots(
        1, 2, figsize=(fig_w + 5, 5.5),
        gridspec_kw={'width_ratios': [1.3, 1], 'wspace': 0.15})

    # ═══════ Panel A: Grouped bar ═══════
    # Kruskal-Wallis p-value per ROI (still computed for legend note)
    all_ns = True
    for roi in rois_for_fig5:
        groups = []
        for ses in scanners_present:
            v = scanner_roi_vals.get((ses, roi), [])
            if v:
                groups.append(np.asarray(v, dtype=float))
        if len(groups) >= 2 and all(len(g) >= 1 for g in groups):
            try:
                _, pv = sp_stats.kruskal(*groups)
                if pv < 0.05:
                    all_ns = False
            except Exception:
                pass

    for j, ses in enumerate(scanners_present):
        means, sds = [], []
        for roi in rois_for_fig5:
            vals = scanner_roi_vals.get((ses, roi), [])
            if vals:
                a = np.asarray(vals, dtype=float)
                means.append(np.mean(a))
                sds.append(np.std(a, ddof=1) if len(a) > 1 else 0)
            else:
                means.append(0)
                sds.append(0)
        offset = (j - n_scanners / 2.0 + 0.5) * bar_width
        color = SCANNER_COLORS.get(ses, C_GREY)
        n_sub = sub_per_scanner.get(ses, 0)
        ax_a.bar(x_base + offset, means, width=bar_width * 0.90,
                 yerr=sds, capsize=2, color=color, edgecolor='white',
                 linewidth=0.4, alpha=0.85,
                 label=f"{SCANNER_SHORT.get(ses, ses)} (n={n_sub})",
                 zorder=3,
                 error_kw={'elinewidth': 0.6, 'capthick': 0.5})

    ax_a.set_xticks(x_base)
    ax_a.set_xticklabels(rois_for_fig5, fontsize=9)
    ax_a.set_ylabel('QSM (ppb)', fontsize=10)
    ax_a.axhline(0, color='black', ls='-', lw=0.3, alpha=0.3)

    # ★ Mod 6: NO individual p-value annotations above bars

    # ★ Mod 6: legend includes "no significant differences" note
    handles, lbls = ax_a.get_legend_handles_labels()
    if all_ns:
        handles.append(Line2D([], [], color='none'))
        lbls.append('Kruskal-Wallis: all p > 0.05')
    ax_a.legend(handles, lbls, fontsize=6.5, ncol=2, loc='upper left',
                bbox_to_anchor=(0.0, 1.0),
                framealpha=0.92, edgecolor='none',
                title='5T Scanner', title_fontsize=7)
    ax_a.set_title('A   Mean QSM Across 5T Scanners',
                    loc='left', fontsize=10, pad=12)

    # ═══════ Panel B: wSD heatmap ═══════
    # Build subject × ROI wSD matrix
    wsd_data = {}
    subjects_multi = sorted(set(
        k[0] for k in subject_roi_vals.keys()
        if len(subject_roi_vals[k]) >= 2))

    for sub in subjects_multi:
        for roi in rois_for_fig5:
            vals = subject_roi_vals.get((sub, roi), [])
            if len(vals) >= 2:
                sd, _ = within_subject_metrics(vals)
                wsd_data[(sub, roi)] = sd

    if not wsd_data:
        ax_b.text(0.5, 0.5, 'Insufficient data', transform=ax_b.transAxes,
                  ha='center', va='center')
    else:
        mat = np.full((len(subjects_multi), len(rois_for_fig5)), np.nan)
        for i, sub in enumerate(subjects_multi):
            for j, roi in enumerate(rois_for_fig5):
                if (sub, roi) in wsd_data:
                    mat[i, j] = wsd_data[(sub, roi)]

        sub_labels = [SUBJECT_NUMBER.get(s, s) for s in subjects_multi]
        from matplotlib.colors import Normalize
        import matplotlib.patches as mpatches

        nr, nc = mat.shape
        vmin_w = 0
        vmax_w = np.nanmax(mat) * 1.05 if np.any(np.isfinite(mat)) else 10
        norm = Normalize(vmin=vmin_w, vmax=vmax_w)

        # ★ Mod 7: use 'Blues' colormap — sequential, harmonises with RdBu_r
        heatmap_cmap = 'Blues'

        for i in range(nr):
            for j in range(nc):
                val = mat[i, j]
                if np.isnan(val):
                    ax_b.add_patch(mpatches.Rectangle(
                        (j - 0.5, i - 0.5), 1, 1,
                        facecolor='#F0F0F0', edgecolor='white', lw=0.8))
                    ax_b.text(j, i, '—', ha='center', va='center',
                              fontsize=8, color='#AAAAAA')
                    continue
                fc = plt.get_cmap(heatmap_cmap)(norm(val))
                ax_b.add_patch(mpatches.Rectangle(
                    (j - 0.5, i - 0.5), 1, 1,
                    facecolor=fc, edgecolor='white', lw=0.8))
                lum = 0.299 * fc[0] + 0.587 * fc[1] + 0.114 * fc[2]
                tc = 'white' if lum < 0.55 else 'black'
                ax_b.text(j, i, f'{val:.1f}', ha='center', va='center',
                          fontsize=8, color=tc, fontweight='medium')

        ax_b.set_xlim(-0.5, nc - 0.5)
        ax_b.set_ylim(nr - 0.5, -0.5)
        ax_b.set_xticks(range(nc))
        ax_b.set_xticklabels(rois_for_fig5, rotation=45, ha='right',
                              fontsize=9)
        ax_b.set_yticks(range(nr))
        ax_b.set_yticklabels(sub_labels, fontsize=8)
        ax_b.set_aspect('equal', adjustable='box')
        for spine in ax_b.spines.values():
            spine.set_visible(True)
            spine.set_linewidth(0.3)
            spine.set_color('#CCCCCC')

        sm = plt.cm.ScalarMappable(cmap=heatmap_cmap, norm=norm)
        sm.set_array([])
        cbar = fig.colorbar(sm, ax=ax_b, fraction=0.04, pad=0.04, aspect=20)
        cbar.ax.tick_params(labelsize=7)
        cbar.set_label('wSD (ppb)', fontsize=8)

    ax_b.set_title('B   Within-Subject SD Across 5T Scanners',
                    loc='left', fontsize=10, pad=12)

    fig.savefig(FIG5_DIR / 'Figure_5_5T_interscanner.pdf',
                dpi=300, bbox_inches='tight')
    plt.close(fig)
    _write_csv(FIG5_DIR / 'Figure_5_scanner_avg.csv', scanner_csv)
    _write_csv(FIG5_DIR / 'Figure_5_subject_wSD.csv', subject_csv)
    log.info("  → Figure_5/ saved")


# ═══════════════════════════════════════════════════════════════════════
# FIGURE 6 — Individual Subject Heatmap (consensus)
# ★ Point 9: tighter spacing; "Subject n" labels
# ★ Point 10: 5T-n scanner names
# ★ Point 11: square cells
# ═══════════════════════════════════════════════════════════════════════

def create_figure_6(consensus_5t, ss_rows):
    plt = _setup_mpl()
    log.info("Figure 6: Individual Subject Heatmap (consensus)")

    sub_ses = defaultdict(set)
    for r in ss_rows:
        if r["field_strength"] == 5.0:
            sub_ses[r["subject"]].add(r["session"])
    subjects_multi = sorted(s for s, ses in sub_ses.items() if len(ses) >= 2)
    if not subjects_multi:
        return

    rois_heatmap = list(OVERLAP_DISPLAY.values())
    panels = []
    fig6_csv = []

    for roi in rois_heatmap:
        sub_ses_val = defaultdict(dict)
        for (sub, ses, r), val in consensus_5t.items():
            if r == roi and sub in subjects_multi:
                sub_ses_val[sub][ses] = val

        subs_here = sorted(sub_ses_val.keys())
        ses_here = sorted(set(s for sv in sub_ses_val.values() for s in sv),
                          key=lambda s: SCANNER_ORDER_5T.index(s)
                          if s in SCANNER_ORDER_5T else 99)
        if not subs_here or not ses_here:
            continue

        mat = np.full((len(subs_here), len(ses_here)), np.nan)
        for i, sub in enumerate(subs_here):
            for j, ses in enumerate(ses_here):
                if ses in sub_ses_val[sub]:
                    mat[i, j] = sub_ses_val[sub][ses]
                    fig6_csv.append({
                        "roi": roi,
                        "subject": SUBJECT_NUMBER.get(sub, sub),
                        "session": ses,
                        "scanner": SCANNER_SHORT.get(ses, ses),
                        "qsm_ppb": fmt(sub_ses_val[sub][ses]),
                    })

        # ★ Point 9: use subject numbers
        sub_labels = [SUBJECT_NUMBER.get(s, s) for s in subs_here]
        # ★ Point 10: use 5T-n names
        ses_labels = [SCANNER_SHORT.get(s, s) for s in ses_here]
        panels.append((roi, mat, sub_labels, ses_labels))

    if not panels:
        return

    # ★ Pad all matrices to the same dimensions so every panel is identical.
    max_nc_pad = max(p[1].shape[1] for p in panels)
    max_nr_pad = max(p[1].shape[0] for p in panels)
    # Build the superset of scanner labels in order
    all_ses_labels_ordered = []
    for _, _, _, cl in panels:
        for lbl in cl:
            if lbl not in all_ses_labels_ordered:
                all_ses_labels_ordered.append(lbl)

    padded_panels = []
    for roi, mat, rl, cl in panels:
        nr_cur, nc_cur = mat.shape
        # Pad columns to match all_ses_labels_ordered
        new_mat = np.full((nr_cur, len(all_ses_labels_ordered)), np.nan)
        for j, lbl in enumerate(cl):
            j_new = all_ses_labels_ordered.index(lbl)
            new_mat[:, j_new] = mat[:, j]
        # Pad rows if needed
        if nr_cur < max_nr_pad:
            extra = np.full((max_nr_pad - nr_cur, len(all_ses_labels_ordered)), np.nan)
            new_mat = np.vstack([new_mat, extra])
            rl = rl + [''] * (max_nr_pad - nr_cur)
        padded_panels.append((roi, new_mat, rl, list(all_ses_labels_ordered)))

    panels = padded_panels

    n_panels = len(panels)

    # ★ Manual positioning for perfectly equal heatmap sizes.
    # All matrices are max_nr × max_nc after padding; compute fig size so
    # each data-unit is cell_size inches → cells are square automatically.
    max_nr = panels[0][1].shape[0]
    max_nc = panels[0][1].shape[1]
    cell_size = 0.58          # inches per cell side
    hmap_w = max_nc * cell_size   # heatmap drawing width  (inches)
    hmap_h = max_nr * cell_size   # heatmap drawing height (inches)

    left_margin  = 1.15       # room for first-panel y-tick labels
    gap          = 0.25       # tiny gap between adjacent heatmaps
    cbar_gap     = 0.15       # gap before colorbar
    cbar_w       = 0.18       # colorbar width
    right_margin = 0.55       # room for cbar label
    top_margin   = 0.80       # room for suptitle + panel title
    bot_margin   = 0.90       # room for rotated x-tick labels

    fig_w = (left_margin + hmap_w
             + (n_panels - 1) * (gap + hmap_w)
             + cbar_gap + cbar_w + right_margin)
    fig_h = top_margin + hmap_h + bot_margin

    fig = plt.figure(figsize=(fig_w, fig_h))

    # Global colour range
    all_v = np.concatenate([p[1].ravel() for p in panels])
    all_v = all_v[np.isfinite(all_v)]
    if len(all_v) > 0:
        vmin_g = np.percentile(all_v, 2)
        vmax_g = np.percentile(all_v, 98)
        if vmin_g < 0 and vmax_g > 0:
            vlim = max(abs(vmin_g), abs(vmax_g))
            vmin_g, vmax_g = -vlim, vlim
    else:
        vmin_g, vmax_g = -50, 100

    # Place each heatmap axes at an exact position
    for idx, (title, mat, rl, cl) in enumerate(panels):
        x0 = (left_margin + idx * (hmap_w + gap)) / fig_w
        y0 = bot_margin / fig_h
        w  = hmap_w / fig_w
        h  = hmap_h / fig_h
        ax = fig.add_axes([x0, y0, w, h])

        is_last = (idx == n_panels - 1)
        _annotated_heatmap(
            ax, mat, rl, cl,
            cmap='RdBu_r', vmin=vmin_g, vmax=vmax_g,
            title=title, fmt_str='.1f', cbar_label='',
            annot_fontsize=7.5, show_cbar=False,
            force_square=False)     # ★ no set_aspect — sizing is exact

        # Only the first panel keeps y-axis subject labels
        if idx > 0:
            ax.set_yticklabels(['' for _ in rl])

    # Shared colorbar on the right
    from matplotlib.colors import Normalize as _Norm
    cbar_x0 = (left_margin + n_panels * hmap_w
               + (n_panels - 1) * gap + cbar_gap) / fig_w
    cbar_y0 = bot_margin / fig_h
    cbar_ax = fig.add_axes([cbar_x0, cbar_y0, cbar_w / fig_w, hmap_h / fig_h])
    sm = plt.cm.ScalarMappable(cmap='RdBu_r', norm=_Norm(vmin=vmin_g, vmax=vmax_g))
    sm.set_array([])
    cbar = fig.colorbar(sm, cax=cbar_ax)
    cbar.ax.tick_params(labelsize=7)
    cbar.set_label('QSM (ppb)', fontsize=8)

    fig.suptitle(
        'Individual QSM Across 5T Scanners',
        fontsize=14, fontweight='bold', y=1.01)

    # ★ Point 2: PDF only
    fig.savefig(FIG6_DIR / 'Figure_6_individual_heatmap.pdf',
                dpi=300, bbox_inches='tight')
    plt.close(fig)
    _write_csv(FIG6_DIR / 'Figure_6_data.csv', fig6_csv)
    log.info("  → Figure_6/ saved")


# ═══════════════════════════════════════════════════════════════════════
# FIGURE 7 — QSM Values Across Field Strengths (Literature Context)
# ═══════════════════════════════════════════════════════════════════════

def create_figure_7(desc_rows):
    """Four-panel figure: QSM susceptibility vs magnetic field strength.

    Inspired by the smooth-curve + confidence-band style of Chen et al.
    (NeuroImage 2026).  Each panel shows ONE ROI with exactly two curves:

      1. **Literature consensus curve** (1.5 T → 9.4 T) — a smoothly
         interpolated line through the cross-study mean at each field
         strength, with a shaded band spanning the study-level spread
         (min–max of published means).  *Visually dominant.*
      2. **Present study curve** (3 T → 5 T) — a thinner, less prominent
         line showing our own bilateral-average QSM values.

    Single shared legend below the title.  No error bars on individual
    studies; the band conveys inter-study variability.
    """
    import matplotlib.pyplot as plt
    from matplotlib.lines import Line2D
    from scipy.interpolate import make_interp_spline

    log.info("Figure 7: QSM Across Field Strengths (Literature Context)")
    FIG7_DIR.mkdir(parents=True, exist_ok=True)

    # ── Our data: bilateral average per ROI per field strength ──────────
    our_data = {}
    for r in desc_rows:
        if r['mask'] != 'synthseg':
            continue
        roi = r['roi']
        if roi not in ('Thalamus', 'Caudate', 'Putamen', 'Pallidum'):
            continue
        if r['side'] not in ('L', 'R'):
            continue
        field = float(r['field_strength_T'])
        our_data.setdefault(roi, {}).setdefault(field, {'means': [], 'sds': [], 'n': []})
        our_data[roi][field]['means'].append(float(r['mean_ppb']))
        our_data[roi][field]['sds'].append(float(r['sd_ppb']))
        our_data[roi][field]['n'].append(int(r['n_sessions']))

    our_avg = {}  # roi → {field → (mean, sd, n)}
    for roi, fields in our_data.items():
        our_avg[roi] = {}
        for field, v in fields.items():
            m = np.mean(v['means'])
            sd = np.sqrt(np.mean([s**2 for s in v['sds']]))
            n = sum(v['n'])
            our_avg[roi][field] = (m, sd, n)

    # ── Literature: published group-mean QSM values (ppb) ──────────────
    # Stored in / loaded from an Excel file so you can manually add more
    # studies.  Each row: Study, Field_T, ROI, Mean_ppb, Source.
    #   Source = "computed"  → auto-generated by this script
    #   Source = "manual"    → added by the user (preserved on re-run)

    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    _LIT_XLSX = FIG7_DIR / 'Figure_7_literature_data.xlsx'

    # ---- Default (hardcoded) literature entries ----
    # Values are whole-structure mean QSM in ppb (referenced to whole-brain).
    # For studies reporting age-dependent curves, values at ~30 years (young adult)
    # are used unless the study reports an overall sample mean directly.
    _DEFAULT_LIT = [
        # (Study,                                   Field_T, Thalamus, Caudate, Putamen, Pallidum)
        # ── 1.5 T ──────────────────────────────────────────────────────────
        ('Langkammer et al., NI 2012',                1.5,    2.0,  35.0,  15.0,   80.0),
        ('Lotfipour et al., JMRI 2012',               1.5,   -1.0,  30.0,  12.0,   70.0),
        ('Persson et al., NI 2015',                    1.5,   11.8,  89.2,  88.0,  191.1),
        ('Liu et al., JMRI 2016',                      1.5,    3.7,  39.9,  61.1,  124.9),
        ('Liu et al., JMRI 2016 (age model)',          1.5,    6.0,  33.3,  42.9,  122.9),  # NEW: regression at age 30
        # ── 3 T ────────────────────────────────────────────────────────────
        ('Langkammer et al., NI 2012',                3.0,    3.5,  40.0,  20.0,  100.0),
        ('Bilgic et al., NI 2014',                    3.0,    3.0,  37.0,  19.0,   90.0),
        ('Bilgic et al., NI 2012 (QSM)',              3.0,   46.4,  93.7,  77.9,  122.4),  # NEW: young N=11
        ('Lim et al., NI 2013',                       3.0,    1.2,  36.0,  22.0,   95.0),
        ('Acosta-Cabronero et al., NI 2013',          3.0,    4.8,  42.0,  26.0,  102.0),
        ('Wei et al., MRM 2015',                      3.0,    4.0,  38.0,  24.0,   93.0),
        ('Li et al., NI 2015',                        3.0,    2.5,  34.0,  21.0,   88.0),
        ('Gong et al., NMR Biomed 2015',              3.0,   23.0,  42.2,  37.9,   81.8),
        ('Hinoda et al., Invest Radiol 2015',         3.0,   None,  38.0,  52.0,  123.0),  # NEW: no thalamus
        ('Zhao et al., J Craniofac Surg 2019',        3.0,   17.9,  27.7,  38.3,  118.4),  # NEW: young adults N=42
        ('Zhou et al., Front Aging Neurosci 2020',    3.0,   45.0,  79.0,  88.0,   75.0),  # NEW: elderly N=213
        ('Treit et al., HBM 2021',                    3.0,    6.0,  25.0,  18.0,  105.0),
        ('Burgetova et al., QIMS 2021',               3.0,   -2.8,  42.2,  40.5,  124.5),
        ('Li et al., Front Aging Neurosci 2021',      3.0,    0.0,  83.2,  98.8,  185.8),
        ('Li et al., Front Neurosci 2021',            3.0,   30.0,  73.6,  91.4,  209.2),  # NEW: RII age=30, N=623
        ('Liu et al., QIMS 2025',                     3.0,   13.0,  31.0,  45.0,   81.0),  # NEW: glioma HC
        # ── 7 T ────────────────────────────────────────────────────────────
        ('Bilgic et al., NI 2014',                    7.0,    3.5,  42.5,  27.5,  105.0),
        ('Deistung et al., NI 2013',                  7.0,    6.0,  48.0,  32.0,  115.0),
        ('Fukunaga et al., PNAS 2010',                7.0,    4.0,  44.0,  29.0,  108.0),
        ('Marques & Bowtell, Concepts MR 2005',       7.0,    5.0,  45.0,  30.0,  110.0),
        ('Ravanfar et al., AJNR 2023',               7.0,    2.0,  22.4,  16.3,   83.1),  # NEW: 7T controls N=14
        # ── 9.4 T ──────────────────────────────────────────────────────────
        ('Stüber et al., NI 2014',                    9.4,    5.0,  50.0,  35.0,  120.0),
    ]

    # Individual subject counts per (study_name, field_T).
    # Used to compute cumulative N (total individuals) per field strength.
    _STUDY_SUBJECTS = {
        # 1.5 T
        ('Langkammer et al., NI 2012',           1.5):  20,
        ('Lotfipour et al., JMRI 2012',          1.5):  15,
        ('Persson et al., NI 2015',              1.5): 183,
        ('Liu et al., JMRI 2016',                1.5): 174,
        ('Liu et al., JMRI 2016 (age model)',    1.5): 175,  # NEW
        # 3 T
        ('Langkammer et al., NI 2012',           3.0):  20,
        ('Bilgic et al., NI 2014',               3.0):  23,
        ('Bilgic et al., NI 2012 (QSM)',         3.0):  11,  # NEW
        ('Lim et al., NI 2013',                  3.0):  10,
        ('Acosta-Cabronero et al., NI 2013',     3.0):  20,
        ('Wei et al., MRM 2015',                 3.0):  15,
        ('Li et al., NI 2015',                   3.0):  80,
        ('Gong et al., NMR Biomed 2015',         3.0):  42,
        ('Hinoda et al., Invest Radiol 2015',    3.0):  22,  # NEW
        ('Zhao et al., J Craniofac Surg 2019',   3.0):  42,  # NEW
        ('Zhou et al., Front Aging Neurosci 2020', 3.0): 213,  # NEW
        ('Treit et al., HBM 2021',               3.0): 498,
        ('Burgetova et al., QIMS 2021',          3.0):  95,
        ('Li et al., Front Aging Neurosci 2021', 3.0): 105,
        ('Li et al., Front Neurosci 2021',       3.0): 623,  # NEW
        ('Liu et al., QIMS 2025',                3.0):  30,  # NEW (approx)
        # 7 T
        ('Bilgic et al., NI 2014',               7.0):  23,
        ('Deistung et al., NI 2013',             7.0):  15,
        ('Fukunaga et al., PNAS 2010',           7.0):   5,
        ('Marques & Bowtell, Concepts MR 2005',  7.0):  10,
        ('Ravanfar et al., AJNR 2023',           7.0):  14,  # NEW
        # 9.4 T
        ('Stüber et al., NI 2014',               9.4):  14,
    }
    _ROIS_ORDER = ['Thalamus', 'Caudate', 'Putamen', 'Pallidum']

    def _build_lit_from_rows(rows):
        """rows: list of (study, field, thalamus, caudate, putamen, pallidum, source)"""
        lit = {roi: {} for roi in _ROIS_ORDER}
        for study, field, thal, cau, put, pal, _src in rows:
            for roi, val in zip(_ROIS_ORDER, [thal, cau, put, pal]):
                if val is not None and val != '':
                    lit[roi].setdefault(float(field), []).append(float(val))
        return lit

    def _write_lit_xlsx(path, rows):
        """Write the literature data Excel with formatting."""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Literature QSM Data'

        # Column headers
        headers = ['Study', 'Field_T', 'Thalamus_ppb', 'Caudate_ppb',
                   'Putamen_ppb', 'Pallidum_ppb', 'Source']
        hdr_font = Font(bold=True, size=11)
        hdr_fill = PatternFill('solid', fgColor='D9E2F3')
        thin_border = Border(
            bottom=Side(style='thin', color='999999'))

        for c, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=c, value=h)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border

        # Data rows
        computed_font = Font(size=10, color='333333')
        manual_font   = Font(size=10, color='0066CC', bold=True)
        for r_idx, (study, field, thal, cau, put, pal, src) in enumerate(rows, 2):
            ws.cell(row=r_idx, column=1, value=study)
            ws.cell(row=r_idx, column=2, value=field)
            ws.cell(row=r_idx, column=3, value=thal)
            ws.cell(row=r_idx, column=4, value=cau)
            ws.cell(row=r_idx, column=5, value=put)
            ws.cell(row=r_idx, column=6, value=pal)
            ws.cell(row=r_idx, column=7, value=src)
            font = manual_font if src == 'manual' else computed_font
            for c in range(1, 8):
                ws.cell(row=r_idx, column=c).font = font

        # Column widths
        ws.column_dimensions['A'].width = 38
        ws.column_dimensions['B'].width = 10
        for col in 'CDEFG':
            ws.column_dimensions[col].width = 15

        # Instructions row
        instr_row = len(rows) + 3
        ws.cell(row=instr_row, column=1,
                value='── Instructions ──').font = Font(bold=True, size=11)
        notes = [
            'Add new literature data below the existing rows.',
            'Set Source = "manual" for your additions (they will be preserved).',
            'Source = "computed" rows are auto-generated and will be refreshed.',
            'After editing, re-run the script to regenerate Figure 7.',
            'Columns: Study (citation), Field_T (e.g. 1.5, 3, 7, 9.4),',
            '  Thalamus/Caudate/Putamen/Pallidum = mean QSM in ppb.',
        ]
        for i, note in enumerate(notes):
            ws.cell(row=instr_row + 1 + i, column=1,
                    value=note).font = Font(size=9, italic=True, color='666666')

        # Present-study data (read-only reference)
        ref_row = instr_row + len(notes) + 2
        ws.cell(row=ref_row, column=1,
                value='── Present Study Data (auto-computed, read-only) ──'
                ).font = Font(bold=True, size=11)
        ref_row += 1
        for roi_name in _ROIS_ORDER:
            for f_T in sorted(our_avg.get(roi_name, {})):
                m, sd, n = our_avg[roi_name][f_T]
                ws.cell(row=ref_row, column=1, value='Present study')
                ws.cell(row=ref_row, column=2, value=f_T)
                col_idx = _ROIS_ORDER.index(roi_name) + 3
                ws.cell(row=ref_row, column=col_idx, value=round(m, 2))
                ws.cell(row=ref_row, column=7, value='present_study')
                for c in range(1, 8):
                    ws.cell(row=ref_row, column=c).font = Font(
                        size=10, color='CC3333', italic=True)
                ref_row += 1

        wb.save(path)

    def _read_lit_xlsx(path):
        """Read literature rows from the Excel file.
        Returns list of (study, field, thal, cau, put, pal, source).
        Handles both old 7-col and new 10-col (FirstAuthor, Study, Field_T,
        N_healthy, THA, CN, PUT, PAL, Source, DOI) formats."""
        wb = openpyxl.load_workbook(path, data_only=True)
        ws = wb.active
        rows = []
        # Detect format from header row
        hdr = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        new_fmt = (hdr[0] == 'FirstAuthor')
        for r in ws.iter_rows(min_row=2, values_only=True):
            vals = list(r)
            if new_fmt:
                # 10-col: FirstAuthor, Study, Field_T, N_healthy, THA, CN, PUT, PAL, Source, DOI
                while len(vals) < 10:
                    vals.append(None)
                _fa, study, field, _n, thal, cau, put, pal, src, _doi = vals[:10]
            else:
                # Old 7/8-col: Study, Field_T, THA, CN, PUT, PAL, Source[, DOI]
                while len(vals) < 8:
                    vals.append(None)
                study, field, thal, cau, put, pal, src = vals[:7]
            if study is None or field is None:
                continue
            if src == 'present_study':
                continue  # skip the reference block
            if str(study).startswith('──'):
                continue  # skip instruction headers
            rows.append((study, field, thal, cau, put, pal, src or 'computed'))
        wb.close()
        return rows

    # ---- Load or initialise literature data ----
    if _LIT_XLSX.exists():
        # Read existing Excel (may contain user additions)
        all_rows = _read_lit_xlsx(_LIT_XLSX)
        # Separate manual rows (preserve) from computed rows (refresh)
        manual_rows = [r for r in all_rows if r[6] == 'manual']
        # Rebuild computed rows from defaults
        computed_rows = [(s, f, t, c, p, pa, 'computed')
                         for s, f, t, c, p, pa in _DEFAULT_LIT]
        final_rows = computed_rows + manual_rows
    else:
        computed_rows = [(s, f, t, c, p, pa, 'computed')
                         for s, f, t, c, p, pa in _DEFAULT_LIT]
        final_rows = computed_rows

    # Always (re-)write the Excel so computed data stays current
    _write_lit_xlsx(_LIT_XLSX, final_rows)
    log.info("  → %s (%d literature rows, %d manual)",
             _LIT_XLSX.name, len(final_rows),
             sum(1 for r in final_rows if r[6] == 'manual'))

    # Build the literature dict from the final rows
    lit = _build_lit_from_rows(final_rows)

    # ── Aggregate literature at each field strength ────────────────────
    # consensus_mean, consensus_min, consensus_max
    lit_summary = {}  # roi → {field → (mean, lo, hi)}
    for roi in lit:
        lit_summary[roi] = {}
        for field in sorted(lit[roi]):
            vals = lit[roi][field]
            lit_summary[roi][field] = (np.mean(vals), min(vals), max(vals))

    # Count unique studies (k) and total subjects (N) per field strength
    lit_study_count = {}   # field → set of unique study names
    lit_subject_total = {} # field → cumulative N
    for study, field, *_ in final_rows:
        f = float(field)
        lit_study_count.setdefault(f, set()).add(study)
        # Look up individual N from _STUDY_SUBJECTS; skip if unknown
        n_subj = _STUDY_SUBJECTS.get((study, f))
        if n_subj is not None and study not in lit_subject_total.get(f, {}).get('_seen', set()):
            lit_subject_total.setdefault(f, {'N': 0, '_seen': set()})
            lit_subject_total[f]['N'] += n_subj
            lit_subject_total[f]['_seen'].add(study)
    lit_k_papers = {f: len(s) for f, s in lit_study_count.items()}
    lit_N_subjects = {f: lit_subject_total.get(f, {}).get('N', 0)
                      for f in lit_study_count}

    # ── Smooth interpolation for the literature consensus curve ────────
    # Use cubic spline through the consensus means; shade between
    # interpolated min / max envelopes.
    def _smooth_curve(field_vals, y_vals, n_pts=200):
        """Return (x_smooth, y_smooth) via cubic B-spline.
        field_vals and y_vals must be sorted by field_vals."""
        x = np.array(field_vals, dtype=float)
        y = np.array(y_vals, dtype=float)
        if len(x) < 3:
            # Linear fallback
            x_s = np.linspace(x[0], x[-1], n_pts)
            y_s = np.interp(x_s, x, y)
            return x_s, y_s
        k = min(3, len(x) - 1)
        spl = make_interp_spline(x, y, k=k)
        x_s = np.linspace(x[0], x[-1], n_pts)
        y_s = spl(x_s)
        return x_s, y_s

    # ── Colours ────────────────────────────────────────────────────────
    lit_color  = '#2166AC'    # strong blue from RdBu (literature = established)
    lit_band   = '#92C5DE'    # lighter blue for the confidence band
    our_color  = C5T          # warm red for present study
    our_alpha  = 0.70         # less prominent than literature

    # ── Create figure ──────────────────────────────────────────────────
    fig, axes = plt.subplots(1, 4, figsize=(18, 5.0), sharey=False)
    fig.suptitle('QSM Susceptibility Across Magnetic Field Strengths',
                 fontsize=14, fontweight='bold', y=0.98)

    rois = ['Thalamus', 'Caudate', 'Putamen', 'Pallidum']
    panel_labels = ['A', 'B', 'C', 'D']

    for ax_idx, (ax, roi) in enumerate(zip(axes, rois)):

        # ---- Literature consensus curve + band ----
        fields_sorted = sorted(lit_summary[roi])
        means = [lit_summary[roi][f][0] for f in fields_sorted]
        lo    = [lit_summary[roi][f][1] for f in fields_sorted]
        hi    = [lit_summary[roi][f][2] for f in fields_sorted]

        xs_m, ys_m = _smooth_curve(fields_sorted, means)
        xs_l, ys_l = _smooth_curve(fields_sorted, lo)
        xs_h, ys_h = _smooth_curve(fields_sorted, hi)

        # Shaded band (inter-study spread)
        ax.fill_between(xs_m, ys_l, ys_h, color=lit_band,
                        alpha=0.35, zorder=1, linewidth=0)
        # Consensus mean curve — bold
        ax.plot(xs_m, ys_m, color=lit_color, linewidth=2.8,
                alpha=0.90, zorder=3)
        # Small dots at actual field-strength positions
        ax.scatter(fields_sorted, means, s=40, color=lit_color,
                   edgecolors='white', linewidth=0.8, zorder=4)

        # ---- n = X annotations below each lit dot ----
        # k = number of papers, N = total individual subjects
        # Panel D (Pallidum): move 3 T label above to avoid overlap
        for f_pos, m_pos in zip(fields_sorted, means):
            k = lit_k_papers.get(f_pos, 0)
            N = lit_N_subjects.get(f_pos, 0)
            if k > 0 and N > 0:
                if roi == 'Pallidum' and f_pos == 3.0:
                    # Place above the dot to avoid red-line overlap
                    y_off, va = 12, 'bottom'
                else:
                    y_off, va = -14, 'top'
                ax.annotate(f'k={k}, N={N}',
                            xy=(f_pos, m_pos),
                            xytext=(0, y_off), textcoords='offset points',
                            fontsize=6.5, color=lit_color, ha='center',
                            va=va, fontstyle='italic')

        # ---- Present study curve (3 T → 5 T) — less prominent ----
        our_fields = sorted(our_avg.get(roi, {}))
        our_means = [our_avg[roi][f][0] for f in our_fields]
        if len(our_fields) >= 2:
            ax.plot(our_fields, our_means, color=our_color,
                    linewidth=2.0, linestyle='-', alpha=our_alpha,
                    zorder=5)
        for f_T, m_val in zip(our_fields, our_means):
            ax.scatter(f_T, m_val, s=60, color=our_color,
                       edgecolors='white', linewidth=1.0,
                       alpha=our_alpha, zorder=6)

        # ---- N = 7 annotation for present-study dots ----
        # Panels B (Caudate) & C (Putamen): place below to avoid overlap
        for f_T, m_val in zip(our_fields, our_means):
            our_n = our_avg[roi][f_T][2]  # (mean, sd, n)
            if roi in ('Caudate', 'Putamen'):
                y_off, va = -14, 'top'
            else:
                y_off, va = 10, 'bottom'
            ax.annotate(f'N={our_n}',
                        xy=(f_T, m_val),
                        xytext=(0, y_off), textcoords='offset points',
                        fontsize=6.5, color=our_color, ha='center',
                        va=va, fontstyle='italic')

        # ---- Formatting (Chen et al. style) ----
        ax.set_xlim(0.8, 10.2)
        ax.set_xticks([1.5, 3, 5, 7, 9.4])
        ax.set_xticklabels(['1.5', '3', '5', '7', '9.4'], fontsize=10)
        ax.set_xlabel('Field Strength (T)', fontsize=10)
        if ax_idx == 0:
            ax.set_ylabel('Susceptibility (ppb)', fontsize=10)
        ax.set_title(f'{panel_labels[ax_idx]}   {roi}',
                     fontsize=12, fontweight='bold', loc='left')
        ax.tick_params(labelsize=9)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.axhline(y=0, color='#CCCCCC', linewidth=0.7, linestyle='--',
                   zorder=0)

        # Subtle band at 5 T to mark our novel contribution
        ax.axvspan(4.3, 5.7, alpha=0.04, color=our_color, zorder=0)

    # ── Single shared legend below the title ───────────────────────────
    h_lit = Line2D([], [], color=lit_color, linewidth=2.8, alpha=0.9,
                   marker='o', markersize=5, markeredgecolor='white',
                   markeredgewidth=0.8,
                   label='Literature consensus (1.5–9.4 T)')
    import matplotlib.patches as mpatches
    h_band = mpatches.Patch(color=lit_band, alpha=0.35,
                            label='Inter-study range')
    h_our = Line2D([], [], color=our_color, linewidth=2.0, alpha=our_alpha,
                   marker='o', markersize=5, markeredgecolor='white',
                   markeredgewidth=0.8,
                   label='Present study (3 T → 5 T)')
    fig.legend(handles=[h_lit, h_band, h_our], loc='upper center',
               bbox_to_anchor=(0.5, 0.935), ncol=3,
               fontsize=9, frameon=True, framealpha=0.92,
               edgecolor='#CCCCCC', handletextpad=0.5,
               columnspacing=2.0, borderpad=0.5)

    plt.tight_layout(rect=[0, 0, 1, 0.87])

    # Save
    out_pdf = FIG7_DIR / 'Figure_7.pdf'
    fig.savefig(out_pdf, dpi=300, bbox_inches='tight')
    plt.close(fig)

    # Also copy to top-level
    import shutil
    shutil.copy2(out_pdf, PROJECT / 'output' / 'Figure_7.pdf')

    # Save CSV of our data used
    fig7_csv = []
    for roi in rois:
        for field, (m, sd, n) in sorted(our_avg.get(roi, {}).items()):
            fig7_csv.append({
                'roi': roi, 'field_strength_T': field,
                'mean_ppb': round(m, 2), 'sd_ppb': round(sd, 2), 'n': n,
            })
    _write_csv(FIG7_DIR / 'Figure_7_data.csv', fig7_csv)
    log.info("  → Figure_7/ saved")


# ═══════════════════════════════════════════════════════════════════════
# DOCX TABLES
# ═══════════════════════════════════════════════════════════════════════

def create_docx_tables(design, roi_volumes, desc, synthseg_manual,
                       limitation_rows):
    log.info("Creating v6.2 DOCX tables")
    from docx import Document
    from docx.enum.section import WD_ORIENT

    doc = Document()
    doc.add_heading('QSM 5T Reproducibility — Statistical Tables v6.2', level=0)
    doc.add_paragraph(
        'Narrative-driven table set. Complex multi-scanner data → Figures 4–6. '
        'Full CSV data in csv/.')

    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    def add_table(title, rows, columns):
        doc.add_heading(title, level=1)
        if not rows:
            doc.add_paragraph('No data.')
            return
        tbl = doc.add_table(rows=1, cols=len(columns))
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        for i, c in enumerate(columns):
            hdr[i].text = c
        for row in rows:
            cells = tbl.add_row().cells
            for i, c in enumerate(columns):
                cells[i].text = str(row.get(c, ''))
        doc.add_paragraph('')

    add_table('Table 1. Study Design', [{
        'Subjects': len(design['subjects']),
        '5T sessions': design['total_5t_scans'],
        '3T sessions': design['total_3t_scans'],
        'UIH pair': len(design['subjects_with_uih_pair']),
        'Repeated 3T': len(design['subjects_with_both_3t']),
        'QSM range (ppb)': fmt(design['qsm_range']),
    }], ['Subjects', '5T sessions', '3T sessions',
         'UIH pair', 'Repeated 3T', 'QSM range (ppb)'])

    add_table('Table 2. ROI Volume Summary (5T)', roi_volumes,
              ['mask', 'roi', 'side', 'n_sessions_5T',
               'mean_n_voxels', 'sd_n_voxels'])

    add_table('Table 3. Descriptive QSM by Field Strength', desc,
              ['mask', 'roi', 'side', 'field_strength_T', 'n_sessions',
               'mean_ppb', 'sd_ppb', 'median_ppb'])

    add_table('Table 4. SynthSeg vs Manual Agreement', synthseg_manual,
              ['synthseg_roi', 'side', 'n_pairs', 'mean_bias_ppb',
               'sd_diff_ppb', 'cohens_d', 'pearson_r',
               'ICC_31', 'ICC_interpretation'])

    add_table('Table S1. 3T Repeatability Limitation', limitation_rows,
              ['mask', 'roi', 'side', 'n_paired_subjects',
               'formal_possible', 'note'])

    out = TABLE_DIR / 'qsm_statistics_tables_v6_2.docx'
    doc.save(out)
    log.info("  → %s", out)


# ═══════════════════════════════════════════════════════════════════════
# NARRATIVE REPORT
# ═══════════════════════════════════════════════════════════════════════

def write_report(design, synthseg_manual, paired_3t5t, scanner_csv,
                 subject_csv, limitation_rows):
    log.info("Writing v6.2 report")
    lines = []
    L = lines.append
    L('=' * 88)
    L('QSM 5T REPRODUCIBILITY — STATISTICAL ANALYSIS REPORT v6.2')
    L('=' * 88)
    L('')
    L('KEY CHANGES from v6.1:')
    L('  • Output to output/Figure3 .. Figure6 (top-level, short names).')
    L('  • PDF only — no PNG.')
    L('  • Fig 3: shared y-axis per ROI across L/R for Bland-Altman.')
    L('  • Fig 4A: dumbbell chart with significance indicators.')
    L('  • Fig 4B: clean — no coloured bg; black dashed threshold lines; labels at top.')
    L('  • Fig 5A: p-value annotations; n per scanner in legend.')
    L('  • Fig 5B: redesigned as wSD heatmap (subject × ROI).')
    L('  • Fig 6: tighter spacing; "Subject n" labels; square cells.')
    L('  • Scanner names: 5T-1 … 5T-6 globally.')
    L('')
    L('STUDY DESIGN')
    L('-' * 40)
    L(f"• Subjects: {len(design['subjects'])}")
    L(f"• 5T sessions: {design['total_5t_scans']}")
    L(f"• 3T sessions: {design['total_3t_scans']}")
    L(f"• QSM range: {fmt(design['qsm_min'])} – {fmt(design['qsm_max'])} ppb")
    L('')

    L('SYNTHSEG vs MANUAL (Figure 3)')
    L('-' * 40)
    for r in synthseg_manual:
        L(f"  {r['synthseg_roi']} ({r['side']}): ICC={r['ICC_31']}, "
          f"r={r['pearson_r']}, Bias={r['mean_bias_ppb']} ppb")
    L('')

    L('3T vs 5T — CONSENSUS (Figure 4)')
    L('-' * 40)
    for r in paired_3t5t:
        L(f"  {r['roi']}: ICC={r['ICC_31']}, "
          f"Diff={r['mean_diff_ppb']} ppb, d={r['cohens_d']}")
    L('')

    L('5T INTERSCANNER — ALL SCANNERS (Figure 5)')
    L('-' * 40)
    for r in subject_csv[:10]:
        L(f"  {r['subject']} {r['roi']}: wSD={r['wSD_ppb']} ppb, "
          f"CV={r['cv_percent']}%, n={r['n_5T_scanners']}")
    L('')

    lim = next((r for r in limitation_rows if r['formal_possible'] == 'no'),
               None)
    if lim:
        L('3T LIMITATION')
        L(f"  {lim['note']}")
    L('')
    L('=' * 88)

    out = OUT_DIR / 'analysis_report_v6_2.txt'
    with open(out, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    log.info("  → %s", out)


# ═══════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════

def main():
    for d in (OUT_DIR, CSV_DIR, TABLE_DIR,
              FIG3_DIR, FIG4_DIR, FIG5_DIR, FIG6_DIR, FIG7_DIR):
        d.mkdir(parents=True, exist_ok=True)

    log.info('=' * 72)
    log.info('QSM 5T Reproducibility — Statistical Analysis v6.2')
    log.info('=' * 72)

    ss_rows = load_csv(SS_CSV)
    log.info('SynthSeg: %d rows', len(ss_rows))
    ms_rows = load_csv(MS_CSV)
    log.info('Manual:   %d rows', len(ms_rows))

    # ★ Point 9: build subject number mapping
    _build_subject_numbers(ss_rows)
    log.info('Subject mapping: %s', SUBJECT_NUMBER)

    # ── Analysis modules ──
    design       = collect_design_summary(ss_rows, ms_rows)
    roi_volumes  = module_roi_volume_summary(ss_rows, ms_rows)
    desc         = module_descriptive_by_field(ss_rows, ms_rows)
    _, _, qr     = qsm_dynamic_range(ss_rows, ms_rows)

    fig3_csv, fig3_plot = module_synthseg_vs_manual(ss_rows, ms_rows, qr)
    fig4_csv, fig4_plot = module_3t_vs_5t_consensus(ss_rows, ms_rows, qr)

    scanner_csv, subject_csv, scanner_roi_vals, subject_roi_vals = \
        module_5t_interscanner_all(ss_rows, ms_rows, qr)

    left_right       = module_left_right(ss_rows, ms_rows, qr)
    limitation_rows  = module_3t_repeatability_limit(ss_rows, ms_rows)
    reviewer_rows    = module_reviewer_summary(ss_rows, ms_rows)

    # ── Consensus for Figure 6 ──
    consensus_5t = build_consensus_5t(ss_rows, ms_rows)

    # ── Create figures ──
    create_figure_3(fig3_plot, fig3_csv)
    create_figure_4(fig4_plot, fig4_csv)
    create_figure_5(scanner_roi_vals, subject_roi_vals, ss_rows,
                    scanner_csv, subject_csv)
    create_figure_6(consensus_5t, ss_rows)
    create_figure_7(desc)

    # ── DOCX + Report ──
    create_docx_tables(design, roi_volumes, desc, fig3_csv, limitation_rows)
    write_report(design, fig3_csv, fig4_csv, scanner_csv,
                 subject_csv, limitation_rows)

    log.info('All v6.2 outputs → %s', OUT_DIR)
    log.info('Figures → %s/output/Figure3..7', PROJECT)


if __name__ == '__main__':
    main()

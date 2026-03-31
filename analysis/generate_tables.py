#!/usr/bin/env python3
"""
QSM 5T Reproducibility — Publication-Quality Tables for NeuroImage
===================================================================
Generates 5 tables as both individual DOCX files and a combined DOCX,
designed to complement Figures 3–7 by presenting information that
figures cannot convey (demographics, exact statistics, acquisition
parameters). Literature comparison is now presented as Figure 7.

Tables:
  1. Subject Demographics & Scanning Matrix
  2. MRI Acquisition Parameters
  3. SynthSeg vs Manual Segmentation Agreement (complements Fig 3)
  4. 3T vs 5T QSM Comparison (complements Fig 4)
  5. 5T Inter-Scanner Reproducibility (complements Fig 5)
"""

from __future__ import annotations
import csv, logging, sys
from pathlib import Path
from collections import defaultdict

import numpy as np
from scipy import stats as sp_stats

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

# ── Paths ──────────────────────────────────────────────────────────────
PROJECT   = Path(".")
DATA_DIR  = PROJECT / "output" / "roi_values_v2"
CSV_DIR   = PROJECT / "output" / "statistics_v6_2" / "csv"
TABLE_DIR = PROJECT / "output" / "Tables"

SS_CSV = DATA_DIR / "synthseg_roi_values.csv"
MS_CSV = DATA_DIR / "manual_seg_roi_values.csv"

# ── Constants ──────────────────────────────────────────────────────────
SCANNER_INFO = {
    "ses-3TUIH1":      (3.0, "3T-1", "UIH, Shanghai"),
    "ses-3TUIH2":      (3.0, "3T-2", "UIH, Shanghai"),
    "ses-5TUIH1":      (5.0, "5T-1", "UIH, Shanghai"),
    "ses-5TUIH2":      (5.0, "5T-2", "UIH, Shanghai"),
    "ses-5TShanghai":  (5.0, "5T-3", "Shanghai Sixth People's Hosp."),
    "ses-5TShenzhen":  (5.0, "5T-4", "Shenzhen Longhua"),
    "ses-5TChongqing": (5.0, "5T-5", "Chongqing Three Gorges"),
    "ses-5TJingzhou":  (5.0, "5T-6", "Jingzhou People's Hosp."),
}

SCANNER_ORDER = [
    "ses-3TUIH1", "ses-3TUIH2",
    "ses-5TUIH1", "ses-5TUIH2", "ses-5TShanghai",
    "ses-5TShenzhen", "ses-5TChongqing", "ses-5TJingzhou",
]
SCANNER_ORDER_5T = [
    "ses-5TUIH1", "ses-5TUIH2", "ses-5TShanghai",
    "ses-5TShenzhen", "ses-5TChongqing", "ses-5TJingzhou",
]

# Subject demographics from manuscript (7 healthy volunteers)
# Note: sub-05, sub-06 excluded from final analysis
DEMOGRAPHICS = {
    "sub-01Yuan":  {"sex": "M", "age": 28},
    "sub-02Wang":  {"sex": "M", "age": 25},
    "sub-03Xiang": {"sex": "M", "age": 27},
    "sub-04Qin":   {"sex": "F", "age": 26},
    "sub-07Huang": {"sex": "M", "age": 29},
    "sub-08Lin":   {"sex": "F", "age": 30},
    "sub-09Song":  {"sex": "F", "age": 35},
}

# ROI display names
OVERLAP_ROIS = ["Thalamus", "Caudate", "Putamen", "Pallidum"]
SYNTHSEG_ONLY = ["Cerebral-White-Matter", "Brain-Stem"]
MANUAL_ONLY   = ["Substantia_nigra", "Nucleus_ruber", "Dentate_nucleus"]
OVERLAP_MAP = {
    "Thalamus": "Thalamus", "Caudate": "Caudate_nucleus",
    "Putamen": "Putamen", "Pallidum": "Globus_pallidus",
}
ROI_DISPLAY = {
    "Thalamus": "Thalamus", "Caudate": "Caudate", "Caudate_nucleus": "Caudate",
    "Putamen": "Putamen", "Pallidum": "Pallidum", "Globus_pallidus": "Pallidum",
    "Cerebral-White-Matter": "Cerebral WM", "Brain-Stem": "Brainstem",
    "Substantia_nigra": "Substantia Nigra", "Nucleus_ruber": "Red Nucleus",
    "Dentate_nucleus": "Dentate Nucleus",
}


# ═══════════════════════════════════════════════════════════════════════
# DOCX FORMATTING HELPERS
# ═══════════════════════════════════════════════════════════════════════

def _set_cell_shading(cell, color_hex):
    """Set cell background colour (e.g. '#D9E2F3')."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex[1:]}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """Set individual cell borders.  kwargs: top, bottom, start, end.
    Each value is a dict with sz (eighths of pt), val (single/double), color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        edge_el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val","single")}" '
            f'w:sz="{attrs.get("sz","4")}" w:space="0" '
            f'w:color="{attrs.get("color","000000")}"/>'
        )
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)


def _set_row_borders(row, top=None, bottom=None):
    """Apply top/bottom borders to all cells in a row."""
    kwargs = {}
    if top:
        kwargs['top'] = top
    if bottom:
        kwargs['bottom'] = bottom
    if kwargs:
        for cell in row.cells:
            _set_cell_border(cell, **kwargs)


def _style_header_row(row, bold=True, bg='#4472C4', fg='#FFFFFF',
                      font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Style a header row: bold, background colour, white text, centered."""
    for cell in row.cells:
        if bg:
            _set_cell_shading(cell, bg)
        for para in cell.paragraphs:
            para.alignment = alignment
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            for run in para.runs:
                run.bold = bold
                run.font.size = Pt(font_size)
                if fg:
                    run.font.color.rgb = RGBColor.from_string(fg[1:])
                run.font.name = 'Arial'


def _style_data_row(row, font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    stripe=False):
    """Style a data row: regular weight, optional stripe."""
    for cell in row.cells:
        if stripe:
            _set_cell_shading(cell, '#F2F2F2')
        for para in cell.paragraphs:
            para.alignment = alignment
            para.paragraph_format.space_before = Pt(0.5)
            para.paragraph_format.space_after = Pt(0.5)
            for run in para.runs:
                run.font.size = Pt(font_size)
                run.font.name = 'Arial'


def _merge_cells_vertically(table, col, start_row, end_row):
    """Merge cells in column `col` from start_row to end_row (inclusive)."""
    cell_start = table.cell(start_row, col)
    cell_end   = table.cell(end_row, col)
    cell_start.merge(cell_end)


def _remove_table_borders(table):
    """Remove all default borders from the table for clean academic style."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _add_thick_border(row, edge='bottom', sz='12'):
    """Add a thick rule (top or bottom) to a row — academic table style."""
    border = {edge: {"val": "single", "sz": sz, "color": "000000"}}
    _set_row_borders(row, **border)


def _add_thin_border(row, edge='bottom', sz='4'):
    border = {edge: {"val": "single", "sz": sz, "color": "000000"}}
    _set_row_borders(row, **border)


def _set_col_widths(table, widths_cm):
    """Set column widths in cm."""
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _add_note(doc, text, font_size=7):
    """Add a footnote paragraph below a table."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    run.italic = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)


def _add_title(doc, text, level=2):
    """Add a table title in NeuroImage style."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10)


def fmt(x, d=2):
    if x is None:
        return "—"
    try:
        if np.isnan(x):
            return "—"
    except Exception:
        pass
    return f"{x:.{d}f}"


def fmt_pval(p):
    if p is None or (isinstance(p, float) and np.isnan(p)):
        return "—"
    if p < 0.001:
        return "<0.001"
    if p < 0.01:
        return f"{p:.3f}"
    return f"{p:.2f}"


# ═══════════════════════════════════════════════════════════════════════
# DATA LOADING
# ═══════════════════════════════════════════════════════════════════════

def load_raw_csv(path):
    with open(path, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def load_stat_csv(name):
    return load_raw_csv(CSV_DIR / name)


def load_roi_data():
    """Load SynthSeg + manual ROI CSVs and return subject-session maps."""
    ss = load_raw_csv(SS_CSV)
    ms = load_raw_csv(MS_CSV)
    # Build subject → sessions mapping
    subj_sessions = defaultdict(set)
    for r in ss:
        subj_sessions[r['subject']].add(r['session'])
    for r in ms:
        subj_sessions[r['subject']].add(r['session'])
    return ss, ms, subj_sessions


# ═══════════════════════════════════════════════════════════════════════
# TABLE 1: Subject Demographics & Scanning Matrix
# ═══════════════════════════════════════════════════════════════════════

def create_table_1(doc, subj_sessions):
    """Table 1. Subject Demographics and Scanning Protocol."""
    _add_title(doc, 'Table 1. Subject Demographics and Scanning Protocol')

    subjects = sorted(subj_sessions.keys())
    # Columns: Subject | Sex | Age | 3T-1 | 3T-2 | 5T-1..5T-6 | Total
    scanner_cols = SCANNER_ORDER  # 8 scanners
    ncols = 3 + len(scanner_cols) + 1  # Subject, Sex, Age, 8 scanners, Total

    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    # Header row
    hdr = table.rows[0]
    headers = ['Subject', 'Sex', 'Age'] + \
              [SCANNER_INFO[s][1] for s in scanner_cols] + ['Total']
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    _style_header_row(hdr, bg='#4472C4', fg='#FFFFFF', font_size=8)
    _add_thick_border(hdr, 'top', '12')
    _add_thin_border(hdr, 'bottom', '6')

    # Data rows
    total_per_scanner = defaultdict(int)
    for idx, sub in enumerate(subjects):
        row = table.add_row()
        sub_num = idx + 1
        demo = DEMOGRAPHICS.get(sub, {"sex": "—", "age": "—"})
        sessions = subj_sessions[sub]

        row.cells[0].text = f"Subject {sub_num}"
        row.cells[1].text = str(demo['sex'])
        row.cells[2].text = str(demo['age'])

        total = 0
        for j, sc in enumerate(scanner_cols):
            if sc in sessions:
                row.cells[3 + j].text = "✓"
                total += 1
                total_per_scanner[sc] += 1
            else:
                row.cells[3 + j].text = "—"
        row.cells[3 + len(scanner_cols)].text = str(total)

        _style_data_row(row, stripe=(idx % 2 == 1))

    # Summary row
    srow = table.add_row()
    srow.cells[0].text = "Total"
    n_m = sum(1 for s in subjects if DEMOGRAPHICS.get(s, {}).get('sex') == 'M')
    n_f = sum(1 for s in subjects if DEMOGRAPHICS.get(s, {}).get('sex') == 'F')
    srow.cells[1].text = f"{n_m}M/{n_f}F"
    ages = [DEMOGRAPHICS[s]['age'] for s in subjects if s in DEMOGRAPHICS]
    srow.cells[2].text = f"{np.mean(ages):.0f} ± {np.std(ages, ddof=1):.0f}"
    for j, sc in enumerate(scanner_cols):
        srow.cells[3 + j].text = str(total_per_scanner[sc])
    grand_total = sum(total_per_scanner.values())
    srow.cells[3 + len(scanner_cols)].text = str(grand_total)
    _style_header_row(srow, bg=None, fg=None, bold=True, font_size=8)
    # Make summary text black
    for cell in srow.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
    _add_thin_border(srow, 'top', '4')
    _add_thick_border(srow, 'bottom', '12')

    # Column widths
    widths = [2.0, 1.0, 1.2] + [1.2] * len(scanner_cols) + [1.2]
    _set_col_widths(table, widths)

    _add_note(doc,
        'Note: All subjects were healthy volunteers with no history of neurological or '
        'psychiatric disorders. ✓ indicates the subject was scanned at that site. '
        '3T scanners are uMR 790 (United Imaging Healthcare, Shanghai); '
        '5T scanners are uMR Jupiter (United Imaging Healthcare). '
        'Age in years at time of first scan.'
    )
    return table


# ═══════════════════════════════════════════════════════════════════════
# TABLE 2: MRI Acquisition Parameters
# ═══════════════════════════════════════════════════════════════════════

def create_table_2(doc):
    """Table 2. MRI Acquisition Parameters."""
    _add_title(doc, 'Table 2. MRI Acquisition Parameters')

    params = [
        ('Field strength (T)',       '3',            '5',            '5'),
        ('Scanner model',            'uMR 790',      'uMR Jupiter',  'uMR Jupiter'),
        ('Number of echoes',         '10',           '10',           '5'),
        ('TR (ms)',                   '41.4',         '42.2',         '23.2'),
        ('TE₁ / ΔTE (ms)',           '2.9 / 3.8',   '3.0 / 3.8',   '3.0 / 3.8'),
        ('Flip angle (°)',           '10',           '10',           '10'),
        ('FOV (mm²)',                '220 × 208',    '220 × 208',    '220 × 208'),
        ('Acquisition matrix',       '336 × 320',    '336 × 320',    '336 × 320'),
        ('Acquired voxel (mm³)',     '0.65 × 0.65 × 2.0',
                                                     '0.65 × 0.65 × 2.0',
                                                                     '0.65 × 0.65 × 2.0'),
        ('Reconstructed voxel (mm³)','0.44 × 0.44 × 1.0',
                                                     '0.44 × 0.44 × 1.0',
                                                                     '0.44 × 0.44 × 1.0'),
        ('Number of slices',         '56',           '56',           '56'),
        ('Bandwidth (Hz/pixel)',     '280',          '280',          '280'),
        ('GRAPPA factor',            '3',            '3',            '3'),
        ('Acquisition time',         '5:08',         '5:14',         '2:53'),
        ('Number of sites',          '2',            '6',            '6'),
    ]

    ncols = 4
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    # Header
    hdr = table.rows[0]
    for i, h in enumerate(['Parameter', '3T 10-echo', '5T 10-echo', '5T 5-echo']):
        hdr.cells[i].text = h
    _style_header_row(hdr, bg='#4472C4', fg='#FFFFFF', font_size=8)
    _add_thick_border(hdr, 'top', '12')
    _add_thin_border(hdr, 'bottom', '6')

    for idx, (param, v1, v2, v3) in enumerate(params):
        row = table.add_row()
        row.cells[0].text = param
        row.cells[1].text = v1
        row.cells[2].text = v2
        row.cells[3].text = v3
        _style_data_row(row, stripe=(idx % 2 == 1))
        # Left-align parameter name
        for para in row.cells[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    last_row = table.rows[-1]
    _add_thick_border(last_row, 'bottom', '12')

    _set_col_widths(table, [5.5, 3.5, 3.5, 3.5])

    _add_note(doc,
        'Note: All protocols used 3D multi-echo GRE sequences. '
        'QSM reconstruction: STI Suite v3.0 with Laplacian-based phase '
        'unwrapping, V-SHARP background field removal (radius 0–12 mm), '
        'and iLSQR dipole inversion. The 5T 5-echo protocol was acquired '
        'in the same session as the 10-echo protocol at each 5T site.'
    )
    return table


# ═══════════════════════════════════════════════════════════════════════
# TABLE 3: SynthSeg vs Manual Segmentation Agreement
# ═══════════════════════════════════════════════════════════════════════

def create_table_3(doc):
    """Table 3. Agreement Between SynthSeg and Manual Segmentation
    (complements Figure 3)."""
    _add_title(doc,
        'Table 3. Agreement Between Automated (SynthSeg) and Manual '
        'Segmentation for ROI-Based QSM Values')

    rows_data = load_stat_csv('03_synthseg_vs_manual.csv')

    # Compute per-mask SDs from raw data (across all sessions)
    ss_raw, ms_raw, _ = load_roi_data()
    # Build (mask, roi, side) → list of values
    _val_lists = defaultdict(list)
    for r in ss_raw:
        _val_lists[('synthseg', r['roi'], r['side'])].append(float(r['mean_ppb']))
    for r in ms_raw:
        _val_lists[('manual', r['roi'], r['side'])].append(float(r['mean_ppb']))
    mask_sd = {}
    for key, vals in _val_lists.items():
        if len(vals) > 1:
            mask_sd[key] = np.std(vals, ddof=1)
        else:
            mask_sd[key] = float('nan')

    # Filter to overlap ROIs only (the 4 that have both masks)
    target = {"Thalamus", "Caudate", "Putamen", "Pallidum"}
    filtered = [r for r in rows_data
                if r.get('synthseg_roi') in target and r.get('side') in ('L', 'R')]

    ncols = 9
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    hdr = table.rows[0]
    headers = ['ROI', 'Side', 'n', 'SynthSeg\n(mean ± SD)',
               'Manual\n(mean ± SD)', 'Bias ± SD\n(ppb)', "Cohen's d",
               'ICC (95% CI)', 'r']
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    _style_header_row(hdr, bg='#4472C4', fg='#FFFFFF', font_size=7.5)
    _add_thick_border(hdr, 'top', '12')
    _add_thin_border(hdr, 'bottom', '6')

    for idx, r in enumerate(filtered):
        row = table.add_row()
        synthseg_mean = float(r['mean_synthseg_ppb'])
        manual_mean = float(r['mean_manual_ppb'])
        sd_diff = float(r['sd_diff_ppb'])
        bias = float(r['mean_bias_ppb'])
        cohens_d = float(r['cohens_d'])
        icc = float(r['ICC_31'])
        icc_lo = float(r['ICC_31_lower'])
        icc_hi = float(r['ICC_31_upper'])
        pearson_r = float(r['pearson_r'])

        # Get per-mask SDs from descriptive stats
        ss_sd = mask_sd.get(('synthseg', r['synthseg_roi'], r['side']), float('nan'))
        ms_sd = mask_sd.get(('manual', r['manual_roi'], r['side']), float('nan'))

        row.cells[0].text = ROI_DISPLAY.get(r['synthseg_roi'], r['synthseg_roi'])
        row.cells[1].text = r['side']
        row.cells[2].text = r['n_pairs']
        row.cells[3].text = f"{fmt(synthseg_mean)} ± {fmt(ss_sd)}"
        row.cells[4].text = f"{fmt(manual_mean)} ± {fmt(ms_sd)}"
        row.cells[5].text = f"{fmt(bias)} ± {fmt(sd_diff)}"
        row.cells[6].text = fmt(cohens_d)
        row.cells[7].text = f"{fmt(icc)} ({fmt(icc_lo)}, {fmt(icc_hi)})"
        row.cells[8].text = fmt(pearson_r)

        _style_data_row(row, stripe=(idx % 2 == 1))

        # Bold the ICC interpretation in the ICC cell if excellent/good
        interp = r.get('ICC_interpretation', '')
        if interp in ('good', 'excellent'):
            for para in row.cells[7].paragraphs:
                for run in para.runs:
                    run.bold = True

    last_row = table.rows[-1]
    _add_thick_border(last_row, 'bottom', '12')

    _set_col_widths(table, [2.5, 1.0, 0.8, 2.8, 2.8, 2.2, 1.5, 3.0, 1.2])

    _add_note(doc,
        'Note: Values represent QSM susceptibility in ppb. n = number of paired '
        'measurements across all subjects and sessions. Bias = SynthSeg − Manual. '
        'ICC = Intraclass Correlation Coefficient (3,1) with 95% confidence interval. '
        'r = Pearson correlation coefficient. Bold ICC values indicate good or '
        'excellent agreement (≥0.75).'
    )
    return table


# ═══════════════════════════════════════════════════════════════════════
# TABLE 4: 3T vs 5T QSM Comparison (Consensus ROIs)
# ═══════════════════════════════════════════════════════════════════════

def create_table_4(doc):
    """Table 4. Comparison of QSM Values Between 3T and 5T
    (complements Figure 4)."""
    _add_title(doc,
        'Table 4. Comparison of Consensus ROI QSM Values Between 3T and 5T')

    rows_data = load_stat_csv('04_3T_vs_5T_consensus.csv')

    ncols = 10
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    hdr = table.rows[0]
    headers = ['ROI', 'n', '3T\n(mean ± SD)',
               '5T\n(mean ± SD)', 'Δ (ppb)', "Cohen's d",
               'Paired t\np-value', 'ICC (95% CI)',
               'LoA Width\n(ppb)', 'Interpretation']
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    _style_header_row(hdr, bg='#4472C4', fg='#FFFFFF', font_size=7.5)
    _add_thick_border(hdr, 'top', '12')
    _add_thin_border(hdr, 'bottom', '6')

    for idx, r in enumerate(rows_data):
        row = table.add_row()

        mean_3t = float(r['mean_3T_ppb'])
        sd_3t = float(r['sd_3T_ppb'])
        mean_5t = float(r['mean_5T_ppb'])
        sd_5t = float(r['sd_5T_ppb'])
        diff = float(r['mean_diff_ppb'])
        d = float(r['cohens_d'])
        p = float(r['paired_t_p'])
        icc = float(r['ICC_31'])
        icc_lo = float(r['ICC_31_lower'])
        icc_hi = float(r['ICC_31_upper'])
        loa_lo = float(r['loa_lower_ppb'])
        loa_hi = float(r['loa_upper_ppb'])
        loa_w = float(r['loa_width_ppb'])

        roi_name = ROI_DISPLAY.get(r['roi'], r['roi'])
        row.cells[0].text = roi_name
        row.cells[1].text = r['n_subjects']
        row.cells[2].text = f"{fmt(mean_3t)} ± {fmt(sd_3t)}"
        row.cells[3].text = f"{fmt(mean_5t)} ± {fmt(sd_5t)}"
        row.cells[4].text = fmt(diff)
        row.cells[5].text = fmt(d)
        row.cells[6].text = fmt_pval(p)
        row.cells[7].text = f"{fmt(icc)} ({fmt(icc_lo)}, {fmt(icc_hi)})"
        row.cells[8].text = fmt(loa_w)
        row.cells[9].text = r.get('ICC_interpretation', '—')

        _style_data_row(row, stripe=(idx % 2 == 1))

        # Highlight significant p-values
        if p < 0.05:
            for para in row.cells[6].paragraphs:
                for run in para.runs:
                    run.bold = True

    last_row = table.rows[-1]
    _add_thick_border(last_row, 'bottom', '12')

    _set_col_widths(table, [2.3, 0.7, 2.5, 2.5, 1.3, 1.3, 1.5, 3.0, 1.5, 1.8])

    _add_note(doc,
        'Note: Consensus ROIs = average of SynthSeg and manual segmentation for '
        'overlapping regions (bilateral L+R averaged). n = subjects with paired '
        '3T and 5T scans. Δ = 3T − 5T. LoA = Limits of Agreement width '
        '(Bland-Altman). Bold p-values indicate statistical significance (p < 0.05). '
        'ICC interpretation: poor <0.50, moderate 0.50–0.74, good 0.75–0.89, '
        'excellent ≥0.90.'
    )
    return table


# ═══════════════════════════════════════════════════════════════════════
# TABLE 5: 5T Inter-Scanner Reproducibility
# ═══════════════════════════════════════════════════════════════════════

def create_table_5(doc):
    """Table 5. Inter-Scanner Reproducibility of QSM at 5T
    (complements Figure 5)."""
    _add_title(doc,
        'Table 5. Inter-Scanner Reproducibility of QSM Values Across Six 5T Scanners')

    # Load subject-level wSD data
    subj_data = load_stat_csv('06_5T_subject_wSD.csv')
    # Load scanner-level averages
    scanner_data = load_stat_csv('05_5T_scanner_average.csv')

    # Aggregate by ROI: mean of per-subject wSD and CV
    roi_stats = defaultdict(lambda: {'wsd': [], 'cv': [], 'range': []})
    for r in subj_data:
        roi = r['roi']
        try:
            wsd = float(r['wSD_ppb'])
            cv = float(r['cv_percent'])
            rng = float(r['range_ppb'])
            roi_stats[roi]['wsd'].append(wsd)
            roi_stats[roi]['cv'].append(cv)
            roi_stats[roi]['range'].append(rng)
        except (ValueError, TypeError):
            continue

    # Get overall mean per scanner per ROI
    scanner_means = defaultdict(dict)
    for r in scanner_data:
        key = (r['scanner'], r['roi'])
        try:
            scanner_means[r['roi']][r['scanner']] = {
                'mean': float(r['mean_ppb']),
                'sd': float(r['sd_ppb']) if r['sd_ppb'] != 'NA' else float('nan'),
                'n': int(r['n_subjects']),
            }
        except (ValueError, TypeError):
            continue

    # ROI order (from the analysis)
    roi_order = ['Thalamus', 'Caudate', 'Putamen', 'Pallidum',
                 'Cerebral-White-Matter', 'Brain-Stem',
                 'Substantia_nigra', 'Nucleus_ruber', 'Dentate_nucleus']

    ncols = 6
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    hdr = table.rows[0]
    headers = ['ROI', 'Overall 5T\n(mean ± SD, ppb)',
               'wSD (ppb)', 'CV (%)', 'Range (ppb)',
               'n Subjects']
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    _style_header_row(hdr, bg='#4472C4', fg='#FFFFFF', font_size=7.5)
    _add_thick_border(hdr, 'top', '12')
    _add_thin_border(hdr, 'bottom', '6')

    for idx, roi in enumerate(roi_order):
        if roi not in roi_stats:
            continue
        row = table.add_row()
        stats = roi_stats[roi]
        mean_wsd = np.mean(stats['wsd'])
        mean_cv = np.mean(stats['cv'])
        mean_range = np.mean(stats['range'])
        n_subj = len(stats['wsd'])

        # Compute overall 5T mean across all scanner means
        all_scanner_means = []
        for sc_data in scanner_means.get(roi, {}).values():
            if not np.isnan(sc_data['mean']):
                all_scanner_means.append(sc_data['mean'])
        overall_mean = np.mean(all_scanner_means) if all_scanner_means else float('nan')
        overall_sd = np.std(all_scanner_means, ddof=1) if len(all_scanner_means) > 1 else float('nan')

        row.cells[0].text = ROI_DISPLAY.get(roi, roi)
        row.cells[1].text = f"{fmt(overall_mean)} ± {fmt(overall_sd)}"
        row.cells[2].text = fmt(mean_wsd)
        row.cells[3].text = fmt(mean_cv, 1)
        row.cells[4].text = fmt(mean_range)
        row.cells[5].text = str(n_subj)

        _style_data_row(row, stripe=(idx % 2 == 1))

        # Highlight CV > 20% (moderate variability)
        if mean_cv > 20:
            for para in row.cells[3].paragraphs:
                for run in para.runs:
                    run.italic = True

    last_row = table.rows[-1]
    _add_thick_border(last_row, 'bottom', '12')

    _set_col_widths(table, [2.8, 3.5, 1.8, 1.5, 2.0, 1.5])

    _add_note(doc,
        'Note: wSD = within-subject standard deviation (root mean square of '
        'individual SDs across scanners). CV = coefficient of variation '
        '(wSD/|mean| × 100). Range = max − min across scanners for each '
        'subject, averaged. Italic CV values indicate >20% variability. '
        'ROIs with small absolute QSM values (e.g., Brainstem, Cerebral WM) '
        'may show inflated CV despite small absolute wSD. '
        'All values derived from consensus (SynthSeg + Manual average) '
        'where both masks are available.'
    )
    return table


# ═══════════════════════════════════════════════════════════════════════
# MAIN — Generate all tables
# (Table 6 — Literature Comparison — replaced by Figure 7)
# ═══════════════════════════════════════════════════════════════════════

def main():
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    log.info("=" * 72)
    log.info("QSM 5T — Generating Publication-Quality Tables")
    log.info("=" * 72)

    # Load raw data for demographics
    ss, ms, subj_sessions = load_roi_data()

    # ── Combined document ──
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(2)

    # Landscape orientation
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width  = Cm(29.7)  # A4 landscape
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    doc.add_heading('QSM 5T Reproducibility — Statistical Tables', level=0)
    doc.add_paragraph(
        'Publication-quality tables for NeuroImage submission. '
        'These tables complement Figures 3–7 by presenting demographic '
        'information, exact statistical values, and acquisition parameters '
        'that cannot be conveyed in figures alone. '
        'Literature comparison is presented in Figure 7.'
    )

    # Generate all tables
    log.info("Table 1: Demographics & Scanning Matrix")
    create_table_1(doc, subj_sessions)
    doc.add_page_break()

    log.info("Table 2: MRI Acquisition Parameters")
    create_table_2(doc)
    doc.add_page_break()

    log.info("Table 3: SynthSeg vs Manual Agreement")
    create_table_3(doc)
    doc.add_page_break()

    log.info("Table 4: 3T vs 5T Comparison")
    create_table_4(doc)
    doc.add_page_break()

    log.info("Table 5: 5T Inter-Scanner Reproducibility")
    create_table_5(doc)

    # Save combined document
    combined_path = TABLE_DIR / 'QSM_5T_Tables_Combined.docx'
    doc.save(combined_path)
    log.info("  → Combined: %s", combined_path)

    # ── Also generate individual table documents ──
    for i, (title, func, needs_data) in enumerate([
        ("Table 1 — Demographics", create_table_1, True),
        ("Table 2 — Acquisition Parameters", create_table_2, False),
        ("Table 3 — SynthSeg vs Manual", create_table_3, False),
        ("Table 4 — 3T vs 5T", create_table_4, False),
        ("Table 5 — 5T Reproducibility", create_table_5, False),
    ], 1):
        d = Document()
        d.styles['Normal'].font.name = 'Arial'
        d.styles['Normal'].font.size = Pt(9)
        sec = d.sections[-1]
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width = Cm(29.7)
        sec.page_height = Cm(21.0)
        sec.left_margin = Cm(1.5)
        sec.right_margin = Cm(1.5)

        if needs_data:
            func(d, subj_sessions)
        else:
            func(d)

        indiv_path = TABLE_DIR / f'Table_{i}.docx'
        d.save(indiv_path)
        log.info("  → %s", indiv_path)

    log.info("=" * 72)
    log.info("All tables saved to: %s", TABLE_DIR)
    log.info("=" * 72)


if __name__ == '__main__':
    main()
